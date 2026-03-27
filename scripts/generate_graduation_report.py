from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

base_dir = r"D:\laragon\www\DATN\website-to-find-workers"
out_path = os.path.join(base_dir, "output", "doc", "bao-cao-do-an-tot-nghiep-website-to-find-workers.docx")
img_physical = os.path.join(base_dir, "output", "find_workers_erd_physical.png")
img_erd = os.path.join(base_dir, "output", "find_workers_erd.png")


doc = Document()

# Default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)

# Title page
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BÁO CÁO ĐỒ ÁN TỐT NGHIỆP")
run.bold = True
run.font.size = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Đề tài: Website-to-find-workers\n(Hệ thống kết nối khách hàng với thợ dịch vụ)")
run.bold = True
run.font.size = Pt(15)

doc.add_paragraph("\n")

meta = [
    "Công nghệ chính: Laravel 12, PHP 8.2, MySQL, Laravel Sanctum, Vite + TailwindCSS",
    "Mô hình người dùng: Admin - Customer - Worker",
    f"Thời điểm tổng hợp báo cáo: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
]
for m in meta:
    doc.add_paragraph(m, style='List Bullet')

doc.add_page_break()

# Section 1
h = doc.add_heading("1. Mục tiêu đề tài", level=1)
doc.add_paragraph(
    "Xây dựng nền tảng web giúp khách hàng tìm thợ theo nhu cầu dịch vụ, "
    "đặt lịch sửa chữa/bảo trì, theo dõi tiến độ công việc, thanh toán trực tuyến "
    "và đánh giá chất lượng sau khi hoàn tất."
)

# Section 2
h = doc.add_heading("2. Kiến trúc và công nghệ triển khai", level=1)
tech_items = [
    "Backend: Laravel 12 (REST API, middleware phân quyền, request validation).",
    "Authentication: Laravel Sanctum + OTP xác thực đăng ký/đăng nhập.",
    "Frontend: Blade + Vite + TailwindCSS.",
    "Database: MySQL với migration quản lý version schema.",
    "Media: Cloudinary (upload ảnh/video mô tả sự cố và kết quả xử lý).",
    "Thanh toán: Tích hợp VNPay, MoMo, ZaloPay (return + IPN webhook).",
    "AI Chatbot: OpenAI + dịch vụ gợi ý thợ, gợi ý ca tương tự, gợi ý video YouTube."
]
for item in tech_items:
    doc.add_paragraph(item, style='List Bullet')

# Section 3
h = doc.add_heading("3. Chức năng đã thực hiện", level=1)

doc.add_heading("3.1. Khối tài khoản và xác thực", level=2)
for item in [
    "Đăng ký, đăng nhập, xác thực OTP, gửi lại OTP.",
    "Phân vai trò tài khoản: admin, customer, worker.",
    "Cập nhật hồ sơ người dùng: avatar, địa chỉ, đổi mật khẩu.",
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("3.2. Khối khách hàng (Customer)", level=2)
for item in [
    "Xem danh mục dịch vụ và hồ sơ thợ công khai.",
    "Đặt lịch sửa chữa tại nhà hoặc tại cửa hàng.",
    "Đính kèm ảnh/video mô tả sự cố khi tạo đơn.",
    "Theo dõi trạng thái đơn, hủy đơn theo điều kiện trạng thái.",
    "Thanh toán COD/chuyển khoản và xem lịch sử đơn của mình.",
    "Đánh giá thợ sau khi hoàn tất công việc."
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("3.3. Khối thợ (Worker)", level=2)
for item in [
    "Quản lý hồ sơ thợ (CCCD, kinh nghiệm, chứng chỉ, khu vực phục vụ).",
    "Nhận việc từ danh sách đơn khả dụng theo chuyên môn dịch vụ.",
    "Cập nhật trạng thái quy trình xử lý đơn (xác nhận, đang làm, chờ hoàn thành...).",
    "Cập nhật chi phí: tiền công, phí linh kiện, phí thuê xe chở (nếu có).",
    "Gửi yêu cầu thanh toán cho khách và xác nhận thanh toán tiền mặt.",
    "Tải lên ảnh/video kết quả sau sửa chữa."
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("3.4. Khối quản trị (Admin)", level=2)
for item in [
    "Dashboard thống kê người dùng, đơn hàng, doanh thu, hoa hồng hệ thống.",
    "Quản lý người dùng: lọc theo vai trò, khóa/mở khóa tài khoản.",
    "Duyệt hồ sơ thợ: chờ duyệt / đã duyệt / từ chối.",
    "Quản lý danh mục dịch vụ: thêm, sửa, ẩn dịch vụ.",
    "Quản lý tất cả đơn đặt lịch trên hệ thống.",
    "Cấu hình 'Assistant Soul' cho chatbot AI từ trang quản trị."
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading("3.5. Chatbot AI hỗ trợ người dùng", level=2)
for item in [
    "Lưu lịch sử hội thoại theo user hoặc guest_token.",
    "Nhận diện ý định tìm thợ theo dịch vụ.",
    "Gợi ý thợ phù hợp theo chuyên môn và dữ liệu ca tương tự.",
    "Xử lý luồng tình huống khẩn cấp bằng bộ từ khóa và mẫu phản hồi cấu hình được.",
    "Hỗ trợ gợi ý video YouTube liên quan lỗi/sự cố."
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 4
doc.add_heading("4. Thiết kế cơ sở dữ liệu", level=1)
doc.add_paragraph(
    "Mô hình dữ liệu được triển khai theo hướng tách rõ user, hồ sơ thợ, dịch vụ, đơn đặt lịch và thanh toán. "
    "ERD bên dưới tổng hợp từ schema thực tế và DBML Sếp cung cấp."
)

table_names = [
    "users", "danh_muc_dich_vu", "ho_so_tho", "tho_dich_vu", "don_dat_lich", "don_dat_lich_dich_vu",
    "danh_gia", "thanh_toan", "otp_codes", "chat_messages", "chat_magic", "app_settings"
]

t = doc.add_table(rows=1, cols=2)
header_cells = t.rows[0].cells
header_cells[0].text = "Bảng"
header_cells[1].text = "Vai trò"

roles = {
    "users": "Quản lý thông tin tài khoản và phân quyền.",
    "danh_muc_dich_vu": "Danh mục dịch vụ hệ thống cung cấp.",
    "ho_so_tho": "Hồ sơ chuyên môn, định vị, trạng thái duyệt của thợ.",
    "tho_dich_vu": "Liên kết nhiều-nhiều giữa thợ và dịch vụ.",
    "don_dat_lich": "Đơn đặt lịch, trạng thái xử lý, chi phí, media.",
    "don_dat_lich_dich_vu": "Liên kết nhiều dịch vụ trong một đơn đặt lịch.",
    "danh_gia": "Đánh giá của khách hàng dành cho thợ.",
    "thanh_toan": "Lưu giao dịch thanh toán và trạng thái.",
    "otp_codes": "Lưu OTP phục vụ xác thực tài khoản.",
    "chat_messages": "Lưu thông điệp chat theo luồng nghiệp vụ.",
    "chat_magic": "Lưu hội thoại AI chatbot và metadata.",
    "app_settings": "Lưu cấu hình hệ thống (bao gồm cấu hình assistant)."
}

for name in table_names:
    row_cells = t.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = roles.get(name, "")

if os.path.exists(img_erd):
    doc.add_paragraph("\nSơ đồ ERD logic:")
    doc.add_picture(img_erd, width=Inches(6.2))

if os.path.exists(img_physical):
    doc.add_paragraph("\nSơ đồ vật lý (physical schema):")
    doc.add_picture(img_physical, width=Inches(6.2))

# Section 5
doc.add_heading("5. Quy trình nghiệp vụ chính", level=1)
process_steps = [
    "(1) Khách hàng đăng ký/đăng nhập -> chọn dịch vụ -> tạo đơn đặt lịch.",
    "(2) Hệ thống kiểm tra loại lịch, khoảng cách phục vụ, tính phí di chuyển.",
    "(3) Đơn được gán thợ cụ thể hoặc broadcast cho nhóm thợ phù hợp.",
    "(4) Thợ nhận việc, cập nhật trạng thái xử lý, cập nhật chi phí thực tế.",
    "(5) Thợ gửi yêu cầu thanh toán -> khách thanh toán (COD/VNPay/MoMo/ZaloPay).",
    "(6) Hệ thống ghi nhận giao dịch thành công, hoàn tất đơn và cho phép đánh giá."
]
for s in process_steps:
    doc.add_paragraph(s, style='List Number')

# Section 6
doc.add_heading("6. Kết quả đạt được", level=1)
for item in [
    "Hoàn thiện hệ thống đa vai trò với luồng nghiệp vụ tương đối đầy đủ.",
    "Triển khai API backend với 61 endpoint phục vụ web/app.",
    "Đưa vào vận hành luồng thanh toán đa cổng và webhook xác nhận giao dịch.",
    "Xây dựng chatbot AI hỗ trợ tư vấn dịch vụ và đề xuất thợ.",
    "Chuẩn hóa dữ liệu bằng migration, request validation và middleware phân quyền."
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 7
doc.add_heading("7. Hạn chế và hướng phát triển", level=1)
for item in [
    "Bổ sung test tự động (unit/integration) cho các luồng nghiệp vụ quan trọng.",
    "Nâng cấp thuật toán matching thợ theo vị trí thời gian thực và điểm uy tín.",
    "Bổ sung realtime notification qua WebSocket/Pusher.",
    "Mở rộng dashboard BI (doanh thu theo dịch vụ/khu vực/thời gian).",
    "Tăng cường giám sát bảo mật và tối ưu hiệu năng truy vấn cho dữ liệu lớn."
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph("\n--- Hết ---")

os.makedirs(os.path.dirname(out_path), exist_ok=True)
doc.save(out_path)
print(out_path)
