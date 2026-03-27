from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

BASE_DIR = r"D:\laragon\www\DATN\website-to-find-workers"
OUT_PATH = os.path.join(BASE_DIR, "output", "doc", "bao-cao-do-an-tot-nghiep-chuan-nop.docx")

# Ảnh minh họa lấy từ project
APPENDIX_IMAGES = [
    os.path.join(BASE_DIR, "output", "find_workers_erd.png"),
    os.path.join(BASE_DIR, "output", "find_workers_erd_physical.png"),
    os.path.join(BASE_DIR, "resources", "views", "customer", "home.blade.php"),  # chỉ để check tồn tại project
]

UI_IMAGE_CANDIDATES = [
    os.path.join(BASE_DIR, "public", "assets", "images", "banner.png"),
    os.path.join(BASE_DIR, "public", "assets", "images", "customer.png"),
    os.path.join(BASE_DIR, "public", "assets", "images", "worker2.png"),
    os.path.join(BASE_DIR, "public", "assets", "images", "nhansu.png"),
]


def set_default_style(document: Document):
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)


def add_center_text(document: Document, text: str, bold=False, size=13, uppercase=False, spacing_after=6):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.upper() if uppercase else text)
    run.bold = bold
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(spacing_after)
    return p


def add_heading_like(document: Document, text: str, level=1):
    return document.add_heading(text, level=level)


def add_toc(paragraph):
    # Field code TOC để Word tự generate mục lục khi Update Field
    run = paragraph.add_run()
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_char_separate = OxmlElement('w:fldChar')
    fld_char_separate.set(qn('w:fldCharType'), 'separate')

    default_text = OxmlElement('w:t')
    default_text.text = 'Nhấn chuột phải vào mục lục và chọn Update Field để cập nhật tự động.'

    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(default_text)
    run._r.append(fld_char_end)


def add_bullets(document: Document, items):
    for item in items:
        document.add_paragraph(item, style='List Bullet')


def add_numbered(document: Document, items):
    for item in items:
        document.add_paragraph(item, style='List Number')


doc = Document()
set_default_style(doc)

# ==========================
# BÌA CHÍNH
# ==========================
add_center_text(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", bold=True, size=14, uppercase=True)
add_center_text(doc, "TRƯỜNG ĐẠI HỌC NHA TRANG", bold=True, size=14, uppercase=True)

doc.add_paragraph("\n\n")
add_center_text(doc, "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", bold=True, size=18, uppercase=True)
add_center_text(doc, "NGÀNH: KỸ THUẬT PHẦN MỀM", bold=True, size=14, uppercase=True)

doc.add_paragraph("\n")
add_center_text(doc, "ĐỀ TÀI", bold=True, size=14, uppercase=True)
add_center_text(doc, "XÂY DỰNG WEBSITE KẾT NỐI KHÁCH HÀNG VỚI THỢ DỊCH VỤ", bold=True, size=15, uppercase=True)
add_center_text(doc, "(website-to-find-workers)", bold=False, size=13)

doc.add_paragraph("\n\n")
add_center_text(doc, "Sinh viên thực hiện: ........................................", size=13)
add_center_text(doc, "MSSV: ............................................................", size=13)
add_center_text(doc, "Giảng viên hướng dẫn: ....................................", size=13)

doc.add_paragraph("\n\n")
add_center_text(doc, f"Khánh Hòa, {datetime.now().strftime('%m/%Y')}", size=13)

doc.add_page_break()

# ==========================
# LỜI CẢM ƠN
# ==========================
add_heading_like(doc, "LỜI CẢM ƠN", level=1)
doc.add_paragraph(
    "Em xin chân thành cảm ơn Quý thầy cô Trường Đại học Nha Trang đã truyền đạt kiến thức và "
    "định hướng trong suốt quá trình học tập. Em đặc biệt cảm ơn giảng viên hướng dẫn đã góp ý, "
    "định hướng chuyên môn để em hoàn thiện đề tài."
)
doc.add_paragraph(
    "Báo cáo tập trung trình bày quá trình phân tích, thiết kế và triển khai hệ thống website kết nối "
    "khách hàng với thợ dịch vụ. Do giới hạn thời gian, báo cáo vẫn còn thiếu sót và em mong nhận được "
    "các góp ý để tiếp tục hoàn thiện sản phẩm trong tương lai."
)

doc.add_page_break()

# ==========================
# MỤC LỤC
# ==========================
add_heading_like(doc, "MỤC LỤC", level=1)
p_toc = doc.add_paragraph()
add_toc(p_toc)

doc.add_page_break()

# ==========================
# MỞ ĐẦU
# ==========================
add_heading_like(doc, "MỞ ĐẦU", level=1)
doc.add_paragraph(
    "Trong bối cảnh nhu cầu sửa chữa, bảo trì thiết bị gia dụng ngày càng tăng, việc tìm thợ uy tín nhanh chóng "
    "vẫn là bài toán thực tế với nhiều hộ gia đình. Đề tài website-to-find-workers được xây dựng nhằm số hóa quy trình "
    "kết nối khách hàng - thợ, tối ưu khâu đặt lịch, theo dõi xử lý và thanh toán."
)
add_bullets(doc, [
    "Mục tiêu: Xây dựng nền tảng kết nối dịch vụ có thể vận hành thực tế.",
    "Phạm vi: Nền tảng web gồm 3 vai trò Admin, Customer, Worker.",
    "Công nghệ: Laravel 12, MySQL, Sanctum, Vite/Tailwind, Cloudinary, tích hợp cổng thanh toán.",
])

# ==========================
# CHƯƠNG 1
# ==========================
add_heading_like(doc, "CHƯƠNG 1. KHẢO SÁT VÀ PHÂN TÍCH YÊU CẦU", level=1)
add_heading_like(doc, "1.1. Bài toán thực tế", level=2)
doc.add_paragraph(
    "Khách hàng thường gặp khó khăn trong việc tìm thợ phù hợp, không rõ giá và thiếu công cụ theo dõi tiến độ. "
    "Trong khi đó, thợ dịch vụ thiếu kênh nhận việc minh bạch, khó tiếp cận khách hàng mới và khó xây dựng uy tín số."
)

add_heading_like(doc, "1.2. Yêu cầu chức năng", level=2)
add_bullets(doc, [
    "Đăng ký/đăng nhập, xác thực OTP, quản lý hồ sơ cá nhân.",
    "Khách hàng tạo đơn đặt lịch theo dịch vụ và thời gian mong muốn.",
    "Thợ nhận đơn, cập nhật trạng thái và chi phí thực tế.",
    "Hệ thống hỗ trợ thanh toán trực tuyến và tiền mặt.",
    "Quản trị viên quản lý người dùng, dịch vụ, hồ sơ thợ, đơn hàng.",
    "Tích hợp chatbot AI hỗ trợ tư vấn nhanh."
])

add_heading_like(doc, "1.3. Yêu cầu phi chức năng", level=2)
add_bullets(doc, [
    "Bảo mật API bằng token (Laravel Sanctum).",
    "Dữ liệu nhất quán thông qua migration + validation.",
    "Khả năng mở rộng thêm dịch vụ/cổng thanh toán.",
    "Giao diện dễ sử dụng với người dùng phổ thông."
])

# ==========================
# CHƯƠNG 2
# ==========================
add_heading_like(doc, "CHƯƠNG 2. THIẾT KẾ HỆ THỐNG", level=1)
add_heading_like(doc, "2.1. Kiến trúc tổng thể", level=2)
doc.add_paragraph(
    "Hệ thống sử dụng kiến trúc client-server. Frontend Blade giao tiếp với backend REST API Laravel. "
    "Tầng backend xử lý phân quyền, nghiệp vụ đặt lịch, thanh toán và chatbot."
)

add_heading_like(doc, "2.2. Thiết kế cơ sở dữ liệu", level=2)
doc.add_paragraph(
    "CSDL được chuẩn hóa theo nhóm thực thể chính: người dùng, dịch vụ, hồ sơ thợ, đơn đặt lịch, thanh toán, đánh giá và chatbot."
)

table = doc.add_table(rows=1, cols=3)
header = table.rows[0].cells
header[0].text = "Nhóm bảng"
header[1].text = "Tên bảng"
header[2].text = "Mô tả"

db_rows = [
    ("Người dùng", "users", "Thông tin tài khoản, vai trò, trạng thái hoạt động."),
    ("Dịch vụ", "danh_muc_dich_vu", "Danh mục dịch vụ sửa chữa/bảo trì."),
    ("Thợ", "ho_so_tho, tho_dich_vu", "Hồ sơ thợ và liên kết chuyên môn dịch vụ."),
    ("Đặt lịch", "don_dat_lich, don_dat_lich_dich_vu", "Thông tin đơn, lịch hẹn, chi phí, trạng thái."),
    ("Đánh giá", "danh_gia", "Đánh giá chất lượng thợ sau khi hoàn thành."),
    ("Thanh toán", "thanh_toan", "Lưu giao dịch từ COD/VNPay/MoMo/ZaloPay."),
    ("AI Chat", "chat_messages, chat_magic, app_settings", "Lưu lịch sử chat và cấu hình assistant."),
]

for r in db_rows:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]

for image_path, caption in [
    (os.path.join(BASE_DIR, "output", "find_workers_erd.png"), "Hình 2.1 - Sơ đồ ERD logic"),
    (os.path.join(BASE_DIR, "output", "find_workers_erd_physical.png"), "Hình 2.2 - Sơ đồ vật lý cơ sở dữ liệu"),
]:
    if os.path.exists(image_path):
        doc.add_paragraph("\n")
        doc.add_picture(image_path, width=Inches(6.2))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ==========================
# CHƯƠNG 3
# ==========================
add_heading_like(doc, "CHƯƠNG 3. TRIỂN KHAI VÀ XÂY DỰNG CHỨC NĂNG", level=1)

add_heading_like(doc, "3.1. Nhóm chức năng xác thực và người dùng", level=2)
add_bullets(doc, [
    "API đăng ký, đăng nhập, xác thực OTP, resend OTP.",
    "Cập nhật thông tin cá nhân, avatar, địa chỉ, mật khẩu.",
    "Phân quyền truy cập theo role admin/customer/worker."
])

add_heading_like(doc, "3.2. Nhóm chức năng đặt lịch và xử lý đơn", level=2)
add_bullets(doc, [
    "Khách tạo đơn theo loại at_home/at_store, chọn khung giờ.",
    "Tính khoảng cách và phí di chuyển theo tọa độ.",
    "Thợ nhận đơn khả dụng theo dịch vụ chuyên môn.",
    "Cập nhật trạng thái: chờ xác nhận -> đang làm -> chờ thanh toán -> đã xong.",
    "Cập nhật chi phí chi tiết: tiền công, phí linh kiện, phí thuê xe chở."
])

add_heading_like(doc, "3.3. Nhóm chức năng thanh toán", level=2)
add_bullets(doc, [
    "Khởi tạo URL thanh toán cho VNPay, MoMo, ZaloPay.",
    "Xử lý callback return và webhook IPN để xác nhận giao dịch.",
    "Ghi nhận lịch sử thanh toán và cập nhật trạng thái đơn."
])

add_heading_like(doc, "3.4. Nhóm chức năng quản trị", level=2)
add_bullets(doc, [
    "Dashboard thống kê người dùng, đơn hàng, doanh thu, hoa hồng hệ thống.",
    "Quản lý người dùng, khóa/mở khóa tài khoản.",
    "Duyệt hồ sơ thợ và điều chỉnh trạng thái hoạt động.",
    "Quản lý danh mục dịch vụ và cấu hình Assistant Soul."
])

add_heading_like(doc, "3.5. Nhóm chức năng chatbot AI", level=2)
add_bullets(doc, [
    "Lưu hội thoại theo user hoặc guest token.",
    "Phân tích ý định tìm dịch vụ, gợi ý thợ phù hợp.",
    "Tạo phản hồi AI dựa trên ngữ cảnh hội thoại + dữ liệu liên quan.",
    "Xử lý tình huống khẩn cấp bằng từ khóa và kịch bản cấu hình."
])

# ==========================
# CHƯƠNG 4
# ==========================
add_heading_like(doc, "CHƯƠNG 4. KIỂM THỬ VÀ ĐÁNH GIÁ", level=1)
add_heading_like(doc, "4.1. Kết quả kiểm thử chức năng", level=2)
add_bullets(doc, [
    "Đăng ký/đăng nhập/OTP hoạt động theo quy trình.",
    "Luồng tạo đơn - nhận đơn - cập nhật trạng thái hoạt động ổn định.",
    "Luồng thanh toán qua cổng trung gian xử lý đúng callback/IPN.",
    "Đánh giá sau dịch vụ cập nhật dữ liệu hồ sơ thợ."
])

add_heading_like(doc, "4.2. Đánh giá ưu điểm", level=2)
add_bullets(doc, [
    "Luồng nghiệp vụ đầy đủ cho mô hình marketplace dịch vụ.",
    "Cấu trúc backend rõ ràng, dễ mở rộng thêm tính năng.",
    "Có tích hợp AI chatbot tăng trải nghiệm tư vấn ban đầu."
])

add_heading_like(doc, "4.3. Hạn chế hiện tại", level=2)
add_bullets(doc, [
    "Chưa có bộ test tự động đầy đủ cho toàn bộ API.",
    "Chưa tối ưu sâu cho tải cao và giám sát thời gian thực.",
    "Một số giao diện cần hoàn thiện UX để phù hợp người dùng lớn tuổi."
])

# ==========================
# CHƯƠNG 5
# ==========================
add_heading_like(doc, "CHƯƠNG 5. KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN", level=1)
add_heading_like(doc, "5.1. Kết luận", level=2)
doc.add_paragraph(
    "Đề tài đã xây dựng thành công nền tảng web kết nối khách hàng với thợ dịch vụ, đáp ứng các chức năng cốt lõi: "
    "đặt lịch, quản lý đơn, thanh toán và đánh giá. Hệ thống có khả năng triển khai thực tế ở quy mô vừa và nhỏ."
)

add_heading_like(doc, "5.2. Hướng phát triển", level=2)
add_numbered(doc, [
    "Bổ sung hệ thống realtime notification và theo dõi vị trí thợ theo thời gian thực.",
    "Xây dựng bộ test tự động (unit/integration/e2e) cho luồng quan trọng.",
    "Ứng dụng machine learning để cải thiện xếp hạng đề xuất thợ.",
    "Mở rộng mô hình đa tỉnh/thành và đa chi nhánh dịch vụ.",
])

# ==========================
# TÀI LIỆU THAM KHẢO
# ==========================
doc.add_page_break()
add_heading_like(doc, "TÀI LIỆU THAM KHẢO", level=1)
refs = [
    "[1] Laravel Documentation, https://laravel.com/docs",
    "[2] Laravel Sanctum Documentation, https://laravel.com/docs/sanctum",
    "[3] Cloudinary Laravel SDK, https://cloudinary.com/documentation/laravel_integration",
    "[4] VNPay API Reference, https://sandbox.vnpayment.vn/apis/docs",
    "[5] MoMo Payment Gateway Docs, https://developers.momo.vn",
    "[6] ZaloPay API Docs, https://docs.zalopay.vn",
    "[7] OpenAI API Docs, https://platform.openai.com/docs",
]
for r in refs:
    doc.add_paragraph(r)

# ==========================
# PHỤ LỤC
# ==========================
doc.add_page_break()
add_heading_like(doc, "PHỤ LỤC", level=1)
add_heading_like(doc, "Phụ lục A. Ảnh sơ đồ hệ thống", level=2)

for image_path, caption in [
    (os.path.join(BASE_DIR, "output", "find_workers_erd.png"), "Phụ lục A.1 - ERD logic"),
    (os.path.join(BASE_DIR, "output", "find_workers_erd_physical.png"), "Phụ lục A.2 - ERD vật lý"),
]:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(6.2))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_heading_like(doc, "Phụ lục B. Ảnh minh họa giao diện", level=2)
added_ui = 0
for image_path in UI_IMAGE_CANDIDATES:
    if os.path.exists(image_path):
        doc.add_picture(image_path, width=Inches(5.8))
        p = doc.add_paragraph(f"Hình giao diện minh họa: {os.path.basename(image_path)}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        added_ui += 1

if added_ui == 0:
    doc.add_paragraph("(Chưa có ảnh chụp màn hình giao diện thực tế. Có thể bổ sung ảnh screenshot khi demo hệ thống)")

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
