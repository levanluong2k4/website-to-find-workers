from datetime import datetime
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_DIR = r"D:\laragon\www\DATN\website-to-find-workers"
OUT_PATH = os.path.join(BASE_DIR, "output", "doc", "bao-cao-csdl-theo-mau-main-tables-v2.docx")
IMG_ERD = os.path.join(BASE_DIR, "output", "find_workers_erd.png")
IMG_PHYSICAL = os.path.join(BASE_DIR, "output", "find_workers_erd_physical.png")

# Dữ liệu bảng chính theo DBML user cung cấp
TABLES = [
    {
        "heading": "3.1. Bang users",
        "function": "Chuc nang: Bang nguoi dung trung tam, luu admin, khach hang va tho.",
        "pk": "Khoa chinh: id",
        "important": "- id (bigint, PK, increment)\n- name, email(unique), password\n- phone, address, avatar\n- role (enum: admin, customer, worker)\n- is_active\n- created_at, updated_at",
        "relation": "- 1-1 voi ho_so_tho theo user_id (chi ap dung cho worker).\n- 1-n voi don_dat_lich theo khach_hang_id (vai tro customer).\n- 1-n voi don_dat_lich theo tho_id (vai tro worker).\n- n-n voi danh_muc_dich_vu qua tho_dich_vu.",
        "constraint": "- unique(email)",
    },
    {
        "heading": "3.2. Bang danh_muc_dich_vu",
        "function": "Chuc nang: Danh muc cac dich vu duoc cung cap trong he thong.",
        "pk": "Khoa chinh: id",
        "important": "- id (bigint, PK, increment)\n- ten_dich_vu\n- mo_ta\n- hinh_anh\n- trang_thai\n- created_at, updated_at",
        "relation": "- n-n voi users qua tho_dich_vu.\n- n-n voi don_dat_lich qua don_dat_lich_dich_vu.",
        "constraint": "- Khuyen nghi unique(ten_dich_vu) de tranh trung lap ten dich vu.",
    },
    {
        "heading": "3.3. Bang ho_so_tho",
        "function": "Chuc nang: Ho so bo sung cho user co vai tro worker.",
        "pk": "Khoa chinh: id",
        "important": "- user_id (unique, FK -> users.id)\n- cccd (unique)\n- kinh_nghiem, chung_chi, bang_gia_tham_khao\n- vi_do, kinh_do, ban_kinh_phuc_vu\n- trang_thai_duyet\n- dang_hoat_dong, trang_thai_hoat_dong\n- danh_gia_trung_binh, tong_so_danh_gia\n- created_at, updated_at",
        "relation": "- n-1 voi users theo user_id (nghiep vu 1-1 do user_id unique).",
        "constraint": "- unique(user_id)\n- unique(cccd)",
    },
    {
        "heading": "3.4. Bang tho_dich_vu",
        "function": "Chuc nang: Bang trung gian the hien ky nang dich vu cua tung worker.",
        "pk": "Khoa chinh: id",
        "important": "- user_id (FK -> users.id)\n- dich_vu_id (FK -> danh_muc_dich_vu.id)\n- created_at, updated_at",
        "relation": "- n-1 voi users theo user_id.\n- n-1 voi danh_muc_dich_vu theo dich_vu_id.",
        "constraint": "- Khuyen nghi unique(user_id, dich_vu_id).",
    },
    {
        "heading": "3.5. Bang don_dat_lich",
        "function": "Chuc nang: Bang nghiep vu trung tam luu don dat lich, tien do xu ly, chi phi va thanh toan.",
        "pk": "Khoa chinh: id",
        "important": "- khach_hang_id (FK -> users.id), tho_id (FK -> users.id)\n- loai_dat_lich, ngay_hen, khung_gio_hen, thoi_gian_hen\n- dia_chi, vi_do, kinh_do\n- mo_ta_van_de, nguyen_nhan, giai_phap\n- hinh_anh_mo_ta, video_mo_ta, hinh_anh_ket_qua, video_ket_qua\n- khoang_cach, phi_di_lai, phi_linh_kien, tien_cong, tien_thue_xe, tong_tien\n- phuong_thuc_thanh_toan, trang_thai_thanh_toan\n- trang_thai, ly_do_huy, thoi_gian_het_han_nhan\n- created_at, updated_at",
        "relation": "- n-1 voi users theo khach_hang_id.\n- n-1 voi users theo tho_id.\n- n-n voi danh_muc_dich_vu qua don_dat_lich_dich_vu.",
        "constraint": "- enum loai_dat_lich, khung_gio_hen, trang_thai, phuong_thuc_thanh_toan.",
    },
    {
        "heading": "3.6. Bang don_dat_lich_dich_vu",
        "function": "Chuc nang: Bang trung gian luu cac dich vu nam trong mot don dat lich.",
        "pk": "Khoa chinh: id",
        "important": "- don_dat_lich_id (FK -> don_dat_lich.id)\n- dich_vu_id (FK -> danh_muc_dich_vu.id)\n- created_at, updated_at",
        "relation": "- n-1 voi don_dat_lich theo don_dat_lich_id.\n- n-1 voi danh_muc_dich_vu theo dich_vu_id.",
        "constraint": "- unique(don_dat_lich_id, dich_vu_id).",
    },
]


def apply_sample_like_format(doc: Document):
    section = doc.sections[0]
    # giống file mẫu
    section.top_margin = Cm(1.524)
    section.bottom_margin = Cm(1.524)
    section.left_margin = Cm(1.778)
    section.right_margin = Cm(1.778)

    normal = doc.styles['Normal']
    normal.font.name = 'Arial'
    normal.font.size = Pt(10.5)

    h1 = doc.styles['Heading 1']
    h1.font.name = 'Arial'
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(24)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Arial'
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(10)


def add_table_block(doc: Document, info: dict):
    doc.add_heading(info['heading'], level=2)
    doc.add_paragraph(info['function'])
    doc.add_paragraph(info['pk'])
    doc.add_paragraph("")

    t = doc.add_table(rows=4, cols=2)
    t.style = 'Table Grid'
    t.autofit = True
    t.rows[0].cells[0].text = 'Noi dung'
    t.rows[0].cells[1].text = 'Chi tiet'
    t.rows[1].cells[0].text = 'Cot quan trong'
    t.rows[1].cells[1].text = info['important']
    t.rows[2].cells[0].text = 'Quan he'
    t.rows[2].cells[1].text = info['relation']
    t.rows[3].cells[0].text = 'Chi muc/rang buoc'
    t.rows[3].cells[1].text = info['constraint']


doc = Document()
apply_sample_like_format(doc)

# mở đầu bám mẫu
p1 = doc.add_paragraph('MO TA DATABASE VA SO DO ERD')
p1.runs[0].bold = True

doc.add_paragraph(
    f"Du an: Website To Find Workers - Tho Tot NTU\nNgay tao tai lieu: {datetime.now().strftime('%d/%m/%Y')}"
)
doc.add_paragraph(
    'Tai lieu nay tong hop cau truc CSDL cho cac bang nghiep vu chinh co moi quan he truc tiep trong he thong.'
)

# mục 1
doc.add_heading('1. Pham vi va cach doc so do', level=1)
for line in [
    'ERD trong tai lieu tap trung vao cac bang nghiep vu chinh dang duoc su dung.',
    'Cac bang he thong/phu tro khong can thiet da duoc luoc bo theo yeu cau.',
    'Thong tin tung bang duoc trinh bay theo chuc nang, khoa chinh, cot quan trong, quan he va rang buoc.'
]:
    doc.add_paragraph(line, style='List Bullet')

# mục 2
doc.add_heading('2. So do ERD tong quan', level=1)
if os.path.exists(IMG_ERD):
    doc.add_picture(IMG_ERD, width=Inches(6.2))
doc.add_paragraph('Hinh 1. So do ERD tong quan cho cac bang nghiep vu chinh trong he thong.')

if os.path.exists(IMG_PHYSICAL):
    doc.add_paragraph('')
    doc.add_picture(IMG_PHYSICAL, width=Inches(6.2))
doc.add_paragraph('Hinh 2. So do vat ly (physical) cua cac bang nghiep vu chinh.')

# mục 3
doc.add_heading('3. Nhom bang nghiep vu chinh', level=1)
for tb in TABLES:
    add_table_block(doc, tb)

# mục 4
doc.add_heading('4. Tong hop quan he database', level=1)
for rel in [
    'users 1-1 ho_so_tho (ap dung cho worker do user_id unique).',
    'users n-n danh_muc_dich_vu qua tho_dich_vu.',
    'users(customer) 1-n don_dat_lich theo khach_hang_id.',
    'users(worker) 1-n don_dat_lich theo tho_id.',
    'don_dat_lich n-n danh_muc_dich_vu qua don_dat_lich_dich_vu.',
]:
    doc.add_paragraph(rel, style='List Bullet')

# mục 5
doc.add_heading('5. Quy tac nghiep vu quan trong', level=1)
for rule in [
    'Moi worker chi co 1 ho so_tho.',
    'Mot don_dat_lich co the gom nhieu dich vu.',
    'Trang thai don_dat_lich phai tuan thu dung luong nghiep vu.',
    'Tong_tien duoc tinh tu cac thanh phan chi phi trong don_dat_lich.',
]:
    doc.add_paragraph(rule, style='List Bullet')

# mục 6
doc.add_heading('6. Nguon doi chieu', level=1)
doc.add_paragraph('- DBML user cung cap (chi lay bang chinh co quan he).')
doc.add_paragraph('- Migration/model trong project website-to-find-workers.')

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
