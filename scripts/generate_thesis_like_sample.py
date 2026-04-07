from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

BASE_DIR = r"D:\laragon\www\DATN\website-to-find-workers"
OUT_PATH = os.path.join(BASE_DIR, "output", "doc", "bao-cao-csdl-theo-mau-main-tables.docx")

IMG_ERD = os.path.join(BASE_DIR, "output", "find_workers_erd.png")
IMG_PHYSICAL = os.path.join(BASE_DIR, "output", "find_workers_erd_physical.png")


doc = Document()

# ===== Header giống file mẫu =====
doc.add_paragraph("MO TA DATABASE VA SO DO ERD")
doc.add_paragraph(
    f"Du an: Website To Find Workers - Tho Tot NTU\nNgay tao tai lieu: {datetime.now().strftime('%d/%m/%Y')}"
)
doc.add_paragraph(
    "Tai lieu nay tong hop cac bang chinh co moi quan he truc tiep trong nghiep vu dat lich tim tho. "
    "Cac bang he thong/phu tro da duoc luoc bo theo yeu cau."
)

# ===== 1 =====
doc.add_heading("1. Pham vi va cach doc so do", level=1)
for b in [
    "Chi mo ta cac bang nghiep vu chinh co quan he trong he thong.",
    "Bo qua cac bang he thong/ho tro khong can thiet cho phan trinh bay do an.",
    "Moi bang duoc mo ta theo: chuc nang, khoa chinh, cot quan trong, quan he va rang buoc.",
]:
    doc.add_paragraph(b, style='List Bullet')

# ===== 2 =====
doc.add_heading("2. So do ERD tong quan", level=1)
if os.path.exists(IMG_ERD):
    doc.add_picture(IMG_ERD, width=Inches(6.2))
cap = doc.add_paragraph("Hinh 1. So do ERD tong quan cho cac bang nghiep vu chinh trong he thong.")
cap.alignment = WD_ALIGN_PARAGRAPH.LEFT

if os.path.exists(IMG_PHYSICAL):
    doc.add_paragraph("")
    doc.add_picture(IMG_PHYSICAL, width=Inches(6.2))
cap2 = doc.add_paragraph("Hinh 2. So do vat ly (physical) cua cac bang chinh.")
cap2.alignment = WD_ALIGN_PARAGRAPH.LEFT

# ===== 3 =====
doc.add_heading("3. Nhom bang nghiep vu chinh", level=1)

TABLES = [
    {
        "heading": "3.1. Bang users",
        "function": "Chuc nang: Bang nguoi dung trung tam, luu admin, khach hang va tho.",
        "pk": "Khoa chinh: id",
        "important": "- name, email, password\n- phone, address, avatar\n- role (admin/customer/worker)\n- is_active",
        "relation": "- 1-1 voi ho_so_tho theo user_id (chi ap dung cho tho).\n- 1-n voi don_dat_lich theo khach_hang_id.\n- 1-n voi don_dat_lich theo tho_id.\n- n-n voi danh_muc_dich_vu qua tho_dich_vu.",
        "constraint": "- unique(email)",
    },
    {
        "heading": "3.2. Bang ho_so_tho",
        "function": "Chuc nang: Ho so bo sung cho user co vai tro tho.",
        "pk": "Khoa chinh: id",
        "important": "- user_id\n- cccd\n- kinh_nghiem, chung_chi, bang_gia_tham_khao\n- vi_do, kinh_do\n- trang_thai_duyet, dang_hoat_dong, trang_thai_hoat_dong\n- danh_gia_trung_binh, tong_so_danh_gia",
        "relation": "- n-1 voi users theo user_id.",
        "constraint": "- unique(user_id)\n- unique(cccd)",
    },
    {
        "heading": "3.3. Bang danh_muc_dich_vu",
        "function": "Chuc nang: Danh muc cac dich vu duoc cung cap trong he thong.",
        "pk": "Khoa chinh: id",
        "important": "- ten_dich_vu\n- mo_ta\n- hinh_anh\n- trang_thai",
        "relation": "- n-n voi users qua tho_dich_vu.\n- n-n voi don_dat_lich qua don_dat_lich_dich_vu.",
        "constraint": "- (khuyen nghi) unique(ten_dich_vu)",
    },
    {
        "heading": "3.4. Bang tho_dich_vu",
        "function": "Chuc nang: Bang lien ket ky nang dich vu cua tung tho.",
        "pk": "Khoa chinh: id",
        "important": "- user_id\n- dich_vu_id",
        "relation": "- n-1 voi users theo user_id.\n- n-1 voi danh_muc_dich_vu theo dich_vu_id.",
        "constraint": "- (khuyen nghi) unique(user_id, dich_vu_id)",
    },
    {
        "heading": "3.5. Bang don_dat_lich",
        "function": "Chuc nang: Bang nghiep vu trung tam luu toan bo don dat lich va trang thai xu ly.",
        "pk": "Khoa chinh: id",
        "important": "- khach_hang_id, tho_id\n- loai_dat_lich, ngay_hen, khung_gio_hen, thoi_gian_hen\n- dia_chi, vi_do, kinh_do\n- mo_ta_van_de, giai_phap\n- phi_di_lai, phi_linh_kien, tien_cong, tien_thue_xe, tong_tien\n- phuong_thuc_thanh_toan, trang_thai_thanh_toan\n- hinh_anh_mo_ta, video_mo_ta, hinh_anh_ket_qua, video_ket_qua\n- trang_thai, ly_do_huy, thoi_gian_het_han_nhan",
        "relation": "- n-1 voi users theo khach_hang_id.\n- n-1 voi users theo tho_id.\n- n-n voi danh_muc_dich_vu qua don_dat_lich_dich_vu.",
        "constraint": "- enum loai_dat_lich, khung_gio_hen, trang_thai, phuong_thuc_thanh_toan",
    },
    {
        "heading": "3.6. Bang don_dat_lich_dich_vu",
        "function": "Chuc nang: Bang lien ket dich vu trong tung don dat lich.",
        "pk": "Khoa chinh: id",
        "important": "- don_dat_lich_id\n- dich_vu_id",
        "relation": "- n-1 voi don_dat_lich theo don_dat_lich_id.\n- n-1 voi danh_muc_dich_vu theo dich_vu_id.",
        "constraint": "- unique(don_dat_lich_id, dich_vu_id)",
    },
]

for t in TABLES:
    doc.add_heading(t["heading"], level=2)
    doc.add_paragraph(t["function"])
    doc.add_paragraph(t["pk"])
    doc.add_paragraph("")

    table = doc.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Noi dung"
    table.rows[0].cells[1].text = "Chi tiet"
    table.rows[1].cells[0].text = "Cot quan trong"
    table.rows[1].cells[1].text = t["important"]
    table.rows[2].cells[0].text = "Quan he"
    table.rows[2].cells[1].text = t["relation"]
    table.rows[3].cells[0].text = "Chi muc/rang buoc"
    table.rows[3].cells[1].text = t["constraint"]

# ===== 4 =====
doc.add_heading("4. Tong hop quan he database", level=1)
for b in [
    "users 1-1 ho_so_tho (voi worker)",
    "users n-n danh_muc_dich_vu qua tho_dich_vu",
    "users(customer) 1-n don_dat_lich (khach_hang_id)",
    "users(worker) 1-n don_dat_lich (tho_id)",
    "don_dat_lich n-n danh_muc_dich_vu qua don_dat_lich_dich_vu",
]:
    doc.add_paragraph(b, style='List Bullet')

# ===== 5 =====
doc.add_heading("5. Quy tac nghiep vu quan trong", level=1)
for b in [
    "Moi worker chi co 1 ho so_tho (rang buoc unique user_id).",
    "Mot don_dat_lich co the gom nhieu dich vu qua bang trung gian don_dat_lich_dich_vu.",
    "Trang thai don_dat_lich phai di theo luong nghiep vu, tranh nhay coc trang thai.",
    "Tong_tien duoc tong hop tu phi_di_lai + phi_linh_kien + tien_cong + tien_thue_xe.",
]:
    doc.add_paragraph(b, style='List Bullet')

# ===== 6 =====
doc.add_heading("6. Nguon doi chieu", level=1)
doc.add_paragraph("- DBML do nguoi dung cung cap (cac bang chinh co quan he).")
doc.add_paragraph("- Migration va model thuc te trong du an website-to-find-workers.")

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
