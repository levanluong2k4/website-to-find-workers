from docx import Document
from docx.shared import Pt, Cm
from datetime import datetime
import os

BASE_DIR = r"D:\laragon\www\DATN\website-to-find-workers"
OUT_PATH = os.path.join(BASE_DIR, "output", "doc", "bao-cao-csdl-5cot-cap-nhat.docx")

TABLES = [
    {
        "name": "users",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("name", "varchar(255)", "Chuỗi ký tự", "Họ tên người dùng"),
            ("email", "varchar(255)", "Email hợp lệ", "UNIQUE"),
            ("email_verified_at", "timestamp", "Ngày giờ", "Có thể null"),
            ("password", "varchar(255)", "Chuỗi băm", "Mật khẩu đã mã hóa"),
            ("phone", "varchar(255)", "Số điện thoại", "Có thể null"),
            ("address", "varchar(255)", "Địa chỉ", "Có thể null"),
            ("avatar", "varchar(255)", "URL/đường dẫn ảnh", "Có thể null"),
            ("role", "varchar(50)", "admin/customer/worker", "Phân quyền người dùng"),
            ("is_active", "tinyint", "0 hoặc 1", "Trạng thái tài khoản"),
            ("remember_token", "varchar(100)", "Chuỗi token", "Có thể null"),
            ("created_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
            ("updated_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
        ],
        "relation": "Quan hệ: 1-1 với ho_so_tho; 1-n với don_dat_lich (khach_hang_id, tho_id); n-n với danh_muc_dich_vu qua tho_dich_vu; 1-n với danh_gia và chat_magic."
    },
    {
        "name": "danh_muc_dich_vu",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("ten_dich_vu", "varchar(255)", "Chuỗi ký tự", "Tên dịch vụ"),
            ("mo_ta", "varchar(255)", "Chuỗi ký tự", "Mô tả ngắn"),
            ("hinh_anh", "varchar(255)", "URL/đường dẫn ảnh", "Ảnh đại diện dịch vụ"),
            ("trang_thai", "tinyint", "0 hoặc 1", "Ẩn/hiện dịch vụ"),
            ("created_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
            ("updated_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
        ],
        "relation": "Quan hệ: n-n với users qua tho_dich_vu; n-n với don_dat_lich qua don_dat_lich_dich_vu."
    },
    {
        "name": "ho_so_tho",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("user_id", "bigint", "ID người dùng", "UNIQUE, FK -> users.id (1-1)"),
            ("cccd", "varchar(255)", "Chuỗi số/ký tự", "UNIQUE"),
            ("kinh_nghiem", "text", "Văn bản dài", "Kinh nghiệm làm việc"),
            ("chung_chi", "varchar(255)", "Chuỗi ký tự", "Có thể null"),
            ("bang_gia_tham_khao", "varchar(255)", "Chuỗi ký tự", "Có thể null"),
            ("vi_do", "decimal", "-90..90", "Vĩ độ"),
            ("kinh_do", "decimal", "-180..180", "Kinh độ"),
            ("ban_kinh_phuc_vu", "int", "Số km", "Bán kính nhận việc"),
            ("trang_thai_duyet", "varchar(50)", "cho_duyet/da_duyet/tu_choi", "Trạng thái admin duyệt"),
            ("ghi_chu_admin", "text", "Văn bản dài", "Có thể null"),
            ("dang_hoat_dong", "tinyint", "0 hoặc 1", "Bật/tắt nhận việc"),
            ("trang_thai_hoat_dong", "varchar(50)", "dang_hoat_dong/dang_ban/ngung_hoat_dong/tam_khoa", "Trạng thái vận hành"),
            ("danh_gia_trung_binh", "decimal", "0.00..5.00", "Điểm đánh giá trung bình"),
            ("tong_so_danh_gia", "int", "Số nguyên >=0", "Tổng lượt đánh giá"),
            ("created_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
            ("updated_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
        ],
        "relation": "Quan hệ: 1-1 với users thông qua user_id."
    },
    {
        "name": "tho_dich_vu",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("user_id", "bigint", "ID người dùng", "FK -> users.id"),
            ("dich_vu_id", "bigint", "ID dịch vụ", "FK -> danh_muc_dich_vu.id"),
            ("created_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
            ("updated_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
        ],
        "relation": "Quan hệ: bảng trung gian n-n giữa users (worker) và danh_muc_dich_vu."
    },
    {
        "name": "don_dat_lich",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("khach_hang_id", "bigint", "ID người dùng", "FK -> users.id"),
            ("tho_id", "bigint", "ID người dùng", "FK -> users.id"),
            ("loai_dat_lich", "varchar(50)", "at_home/at_store", "Loại đặt lịch"),
            ("thoi_gian_hen", "datetime", "Ngày giờ", "Thời gian hẹn cụ thể"),
            ("ngay_hen", "date", "Ngày", "Ngày hẹn"),
            ("khung_gio_hen", "varchar(50)", "08:00-10:00/10:00-12:00/12:00-14:00/14:00-17:00", "Khung giờ"),
            ("khoang_cach", "decimal", "Số thực >=0", "Khoảng cách km"),
            ("phi_di_lai", "decimal", "Số tiền >=0", "Phí di chuyển"),
            ("phi_linh_kien", "decimal", "Số tiền >=0", "Phí linh kiện"),
            ("thue_xe_cho", "tinyint", "0 hoặc 1", "Có thuê xe chở"),
            ("tien_thue_xe", "decimal", "Số tiền >=0", "Chi phí thuê xe"),
            ("ghi_chu_linh_kien", "text", "Văn bản dài", "Ghi chú linh kiện"),
            ("dia_chi", "varchar(255)", "Chuỗi ký tự", "Địa chỉ làm dịch vụ"),
            ("vi_do", "decimal", "-90..90", "Vĩ độ"),
            ("kinh_do", "decimal", "-180..180", "Kinh độ"),
            ("mo_ta_van_de", "text", "Văn bản dài", "Mô tả lỗi"),
            ("giai_phap", "text", "Văn bản dài", "Giải pháp"),
            ("hinh_anh_mo_ta", "json", "Mảng URL ảnh", "Ảnh trước sửa"),
            ("video_mo_ta", "varchar(255)", "URL video", "Video mô tả"),
            ("hinh_anh_ket_qua", "json", "Mảng URL ảnh", "Ảnh sau sửa"),
            ("video_ket_qua", "varchar(255)", "URL video", "Video kết quả"),
            ("trang_thai", "varchar(50)", "cho_xac_nhan/da_xac_nhan/dang_lam/cho_hoan_thanh/cho_thanh_toan/da_xong/da_huy", "Trạng thái đơn"),
            ("tien_cong", "decimal", "Số tiền >=0", "Công sửa chữa"),
            ("ly_do_huy", "text", "Văn bản dài", "Lý do hủy"),
            ("tong_tien", "decimal", "Số tiền >=0", "Tổng thanh toán"),
            ("phuong_thuc_thanh_toan", "varchar(50)", "cod/transfer", "Phương thức thanh toán"),
            ("trang_thai_thanh_toan", "tinyint", "0 hoặc 1", "Đã thanh toán/chưa"),
            ("thoi_gian_het_han_nhan", "timestamp", "Ngày giờ", "Hạn nhận đơn"),
            ("created_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
            ("updated_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
        ],
        "relation": "Quan hệ: n-1 với users (khach_hang_id, tho_id); n-n với danh_muc_dich_vu qua don_dat_lich_dich_vu; 1-1 nghiệp vụ với thanh_toan và danh_gia."
    },
    {
        "name": "don_dat_lich_dich_vu",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("don_dat_lich_id", "bigint", "ID đơn đặt lịch", "FK -> don_dat_lich.id"),
            ("dich_vu_id", "bigint", "ID dịch vụ", "FK -> danh_muc_dich_vu.id"),
            ("created_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
            ("updated_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
        ],
        "relation": "Quan hệ: bảng trung gian n-n giữa don_dat_lich và danh_muc_dich_vu."
    },
    {
        "name": "danh_gia",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("don_dat_lich_id", "bigint", "ID đơn đặt lịch", "FK -> don_dat_lich.id"),
            ("so_lan_sua", "int", "0..1", "Số lần sửa đánh giá, tối đa 1 lần"),
            ("nguoi_danh_gia_id", "bigint", "ID người dùng", "FK -> users.id"),
            ("nguoi_bi_danh_gia_id", "bigint", "ID người dùng", "FK -> users.id"),
            ("so_sao", "int", "1..5", "Số sao đánh giá"),
            ("nhan_xet", "text", "Văn bản dài", "Nội dung nhận xét"),
            ("chuyen_mon", "int", "1..5", "Điểm chuyên môn"),
            ("thai_do", "int", "1..5", "Điểm thái độ"),
            ("dung_gio", "int", "1..5", "Điểm đúng giờ"),
            ("gia_ca", "int", "1..5", "Điểm giá cả"),
            ("created_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
            ("updated_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
        ],
        "relation": "Quan hệ: n-1 với don_dat_lich; n-1 với users qua nguoi_danh_gia_id và nguoi_bi_danh_gia_id."
    },
    {
        "name": "thanh_toan",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("don_dat_lich_id", "bigint", "ID đơn đặt lịch", "FK -> don_dat_lich.id, 1-1 nghiệp vụ"),
            ("so_tien", "decimal", "Số tiền >=0", "Tổng tiền thanh toán"),
            ("phuong_thuc", "varchar(50)", "cash/vnpay/momo/zalopay", "Phương thức thanh toán"),
            ("ma_giao_dich", "varchar(255)", "Chuỗi ký tự", "Mã giao dịch cổng thanh toán"),
            ("trang_thai", "varchar(50)", "pending/success/failed", "Trạng thái giao dịch"),
            ("thong_tin_extra", "json", "JSON object", "Raw response từ gateway"),
            ("created_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
            ("updated_at", "timestamp", "Ngày giờ", "Laravel timestamp"),
        ],
        "relation": "Quan hệ: n-1 với don_dat_lich (thiết kế nghiệp vụ là 1-1: một đơn một thanh toán tổng)."
    },
    {
        "name": "chat_magic",
        "rows": [
            ("id", "bigint", "Số nguyên dương", "PK, tự tăng"),
            ("user_id", "bigint", "ID người dùng", "FK -> users.id, có thể null cho khách vãng lai"),
            ("guest_token", "varchar(100)", "Chuỗi token", "Định danh guest"),
            ("sender", "varchar(50)", "user/assistant/system", "Người gửi tin nhắn"),
            ("text", "longtext", "Văn bản rất dài", "Nội dung hội thoại"),
            ("meta", "json", "JSON object", "Dữ liệu phụ trợ"),
            ("created_at", "timestamp", "Ngày giờ", "Thời điểm tạo bản ghi"),
        ],
        "relation": "Quan hệ: n-1 với users qua user_id."
    },
]


def init_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(2)
    sec.bottom_margin = Cm(2)
    sec.left_margin = Cm(2)
    sec.right_margin = Cm(2)
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal.font.size = Pt(13)
    return doc


doc = init_doc()
doc.add_heading('MÔ TẢ CƠ SỞ DỮ LIỆU (BẢNG CHÍNH + QUAN HỆ LIÊN QUAN)', level=1)
doc.add_paragraph(f"Dự án: Website-to-find-workers | Ngày cập nhật: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
doc.add_paragraph('Định dạng bảng: STT | Tên thuộc tính | Kiểu dữ liệu | Miền giá trị | Ghi chú')

for idx, table in enumerate(TABLES, start=1):
    doc.add_heading(f"{idx}. Bảng {table['name']}", level=2)
    t = doc.add_table(rows=1, cols=5)
    t.style = 'Table Grid'
    h = t.rows[0].cells
    h[0].text = 'STT'
    h[1].text = 'Tên thuộc tính'
    h[2].text = 'Kiểu dữ liệu'
    h[3].text = 'Miền giá trị'
    h[4].text = 'Ghi chú'

    for i, row in enumerate(table['rows'], start=1):
        r = t.add_row().cells
        r[0].text = str(i)
        r[1].text = row[0]
        r[2].text = row[1]
        r[3].text = row[2]
        r[4].text = row[3]

    doc.add_paragraph(table['relation'])

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
