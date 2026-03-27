from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os

BASE_DIR = r"D:\laragon\www\DATN\website-to-find-workers"
OUT_PATH = os.path.join(BASE_DIR, "output", "doc", "bao-cao-do-an-tot-nghiep-chuan-nop-v2.docx")

CORE_TABLES = [
    {
        "name": "users",
        "desc": "Lưu thông tin tài khoản hệ thống (admin/customer/worker).",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT"),
            ("name", "varchar(255)", "NOT NULL"),
            ("email", "varchar(255)", "UNIQUE"),
            ("password", "varchar(255)", "NOT NULL"),
            ("phone", "varchar(255)", ""),
            ("address", "varchar(255)", ""),
            ("avatar", "varchar(255)", ""),
            ("role", "varchar(50)", "enum: admin, customer, worker"),
            ("is_active", "tinyint", ""),
            ("created_at", "timestamp", ""),
            ("updated_at", "timestamp", ""),
        ]
    },
    {
        "name": "danh_muc_dich_vu",
        "desc": "Danh mục dịch vụ sửa chữa/bảo trì.",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT"),
            ("ten_dich_vu", "varchar(255)", ""),
            ("mo_ta", "varchar(255)", ""),
            ("hinh_anh", "varchar(255)", ""),
            ("trang_thai", "tinyint", ""),
            ("created_at", "timestamp", ""),
            ("updated_at", "timestamp", ""),
        ]
    },
    {
        "name": "ho_so_tho",
        "desc": "Hồ sơ chuyên môn và trạng thái duyệt của thợ (1-1 với users).",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT"),
            ("user_id", "bigint", "UNIQUE, FK -> users.id"),
            ("cccd", "varchar(255)", "UNIQUE"),
            ("kinh_nghiem", "text", ""),
            ("chung_chi", "varchar(255)", ""),
            ("bang_gia_tham_khao", "varchar(255)", ""),
            ("vi_do", "decimal", ""),
            ("kinh_do", "decimal", ""),
            ("ban_kinh_phuc_vu", "int", ""),
            ("trang_thai_duyet", "varchar(50)", "enum: cho_duyet, da_duyet, tu_choi"),
            ("dang_hoat_dong", "tinyint", ""),
            ("trang_thai_hoat_dong", "varchar(50)", "enum: dang_hoat_dong, dang_ban, ngung_hoat_dong, tam_khoa"),
            ("danh_gia_trung_binh", "decimal", "0.00 -> 5.00"),
            ("tong_so_danh_gia", "int", ""),
            ("created_at", "timestamp", ""),
            ("updated_at", "timestamp", ""),
        ]
    },
    {
        "name": "tho_dich_vu",
        "desc": "Bảng trung gian N-N giữa thợ và dịch vụ.",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT"),
            ("user_id", "bigint", "FK -> users.id"),
            ("dich_vu_id", "bigint", "FK -> danh_muc_dich_vu.id"),
            ("created_at", "timestamp", ""),
            ("updated_at", "timestamp", ""),
        ]
    },
    {
        "name": "don_dat_lich",
        "desc": "Bảng nghiệp vụ chính quản lý đơn đặt lịch và trạng thái xử lý.",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT"),
            ("khach_hang_id", "bigint", "FK -> users.id"),
            ("tho_id", "bigint", "FK -> users.id"),
            ("loai_dat_lich", "varchar(50)", "enum: at_home, at_store"),
            ("ngay_hen", "date", ""),
            ("khung_gio_hen", "varchar(50)", "enum: 08:00-10:00, 10:00-12:00, 12:00-14:00, 14:00-17:00"),
            ("dia_chi", "varchar(255)", ""),
            ("vi_do", "decimal", ""),
            ("kinh_do", "decimal", ""),
            ("mo_ta_van_de", "text", ""),
            ("hinh_anh_mo_ta", "json", ""),
            ("video_mo_ta", "varchar(255)", ""),
            ("trang_thai", "varchar(50)", "enum: cho_xac_nhan, da_xac_nhan, dang_lam, cho_hoan_thanh, cho_thanh_toan, da_xong, da_huy"),
            ("phi_di_lai", "decimal", ""),
            ("phi_linh_kien", "decimal", ""),
            ("tien_cong", "decimal", ""),
            ("tong_tien", "decimal", ""),
            ("phuong_thuc_thanh_toan", "varchar(50)", "enum: cod, transfer"),
            ("trang_thai_thanh_toan", "tinyint", ""),
            ("created_at", "timestamp", ""),
            ("updated_at", "timestamp", ""),
        ]
    },
    {
        "name": "don_dat_lich_dich_vu",
        "desc": "Bảng trung gian N-N giữa đơn đặt lịch và danh mục dịch vụ.",
        "fields": [
            ("id", "bigint", "PK, AUTO_INCREMENT"),
            ("don_dat_lich_id", "bigint", "FK -> don_dat_lich.id"),
            ("dich_vu_id", "bigint", "FK -> danh_muc_dich_vu.id"),
            ("created_at", "timestamp", ""),
            ("updated_at", "timestamp", ""),
        ]
    },
]

RELATIONS = [
    "users (worker) 1 - 1 ho_so_tho",
    "users (worker) N - N danh_muc_dich_vu thông qua tho_dich_vu",
    "users (customer) 1 - N don_dat_lich (khach_hang_id)",
    "users (worker) 1 - N don_dat_lich (tho_id)",
    "don_dat_lich N - N danh_muc_dich_vu thông qua don_dat_lich_dich_vu",
]


def set_default(document: Document):
    style = document.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)


def center(doc: Document, text: str, size=13, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


def add_toc_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement('w:fldChar'); separate.set(qn('w:fldCharType'), 'separate')
    text = OxmlElement('w:t'); text.text = 'Update Field để cập nhật mục lục.'
    end = OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'), 'end')
    run._r.append(begin); run._r.append(instr); run._r.append(separate); run._r.append(text); run._r.append(end)


doc = Document()
set_default(doc)

# Bìa
center(doc, "BỘ GIÁO DỤC VÀ ĐÀO TẠO", 14, True)
center(doc, "TRƯỜNG ĐẠI HỌC NHA TRANG", 14, True)
doc.add_paragraph("\n\n")
center(doc, "BÁO CÁO ĐỒ ÁN TỐT NGHIỆP", 18, True)
center(doc, "ĐỀ TÀI: WEBSITE-TO-FIND-WORKERS", 15, True)
doc.add_paragraph("\n")
center(doc, f"Khánh Hòa, {datetime.now().strftime('%m/%Y')}", 13, False)
doc.add_page_break()

# Lời cảm ơn
doc.add_heading("LỜI CẢM ƠN", level=1)
doc.add_paragraph("Em xin chân thành cảm ơn Quý thầy cô và giảng viên hướng dẫn đã hỗ trợ trong quá trình thực hiện đề tài.")
doc.add_page_break()

# Mục lục
doc.add_heading("MỤC LỤC", level=1)
p = doc.add_paragraph(); add_toc_field(p)
doc.add_page_break()

# Chương 1-5 ngắn gọn
doc.add_heading("CHƯƠNG 1. GIỚI THIỆU ĐỀ TÀI", level=1)
doc.add_paragraph("Hệ thống kết nối khách hàng với thợ dịch vụ, hỗ trợ đặt lịch, theo dõi xử lý, thanh toán và đánh giá.")

doc.add_heading("CHƯƠNG 2. PHÂN TÍCH VÀ THIẾT KẾ CSDL", level=1)
doc.add_heading("2.1. Danh sách bảng chính có quan hệ", level=2)
for t in CORE_TABLES:
    doc.add_paragraph(f"- {t['name']}: {t['desc']}")

doc.add_heading("2.2. Mô tả chi tiết từng bảng", level=2)
for idx, t in enumerate(CORE_TABLES, start=1):
    doc.add_heading(f"Bảng {idx}: {t['name']}", level=3)
    doc.add_paragraph(t['desc'])
    tbl = doc.add_table(rows=1, cols=4)
    h = tbl.rows[0].cells
    h[0].text = "STT"; h[1].text = "Trường"; h[2].text = "Kiểu dữ liệu"; h[3].text = "Ràng buộc/Ghi chú"
    for i, f in enumerate(t['fields'], start=1):
        r = tbl.add_row().cells
        r[0].text = str(i)
        r[1].text = f[0]
        r[2].text = f[1]
        r[3].text = f[2]


doc.add_heading("2.3. Quan hệ giữa các bảng", level=2)
for rel in RELATIONS:
    doc.add_paragraph(rel, style='List Bullet')

for image_path, caption in [
    (os.path.join(BASE_DIR, "output", "find_workers_erd.png"), "Hình 2.1 - ERD logic"),
    (os.path.join(BASE_DIR, "output", "find_workers_erd_physical.png"), "Hình 2.2 - ERD vật lý"),
]:
    if os.path.exists(image_path):
        doc.add_paragraph("\n")
        doc.add_picture(image_path, width=Inches(6.1))
        cp = doc.add_paragraph(caption); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading("CHƯƠNG 3. TRIỂN KHAI CHỨC NĂNG", level=1)
doc.add_paragraph("Triển khai các module: xác thực OTP, quản lý hồ sơ thợ, đặt lịch, xử lý đơn, thanh toán, chatbot AI.")

doc.add_heading("CHƯƠNG 4. KIỂM THỬ VÀ ĐÁNH GIÁ", level=1)
doc.add_paragraph("Kết quả thử nghiệm cho thấy luồng nghiệp vụ chính hoạt động ổn định ở mức đề tài tốt nghiệp.")

doc.add_heading("CHƯƠNG 5. KẾT LUẬN", level=1)
doc.add_paragraph("Đề tài đáp ứng mục tiêu xây dựng nền tảng kết nối dịch vụ với kiến trúc có thể mở rộng.")

doc.add_page_break()
doc.add_heading("TÀI LIỆU THAM KHẢO", level=1)
for ref in [
    "[1] Laravel Documentation - https://laravel.com/docs",
    "[2] MySQL Documentation - https://dev.mysql.com/doc",
    "[3] Laravel Sanctum - https://laravel.com/docs/sanctum",
]:
    doc.add_paragraph(ref)


doc.add_page_break()
doc.add_heading("PHỤ LỤC", level=1)
doc.add_paragraph("Phụ lục A: Sơ đồ ERD logic và vật lý.")

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(OUT_PATH)
