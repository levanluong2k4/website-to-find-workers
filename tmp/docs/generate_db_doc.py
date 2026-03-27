from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"d:\laragon\www\DATN\website-to-find-workers")
OUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs"
DOC_PATH = OUT_DIR / "mo-ta-database-va-erd-thototntu.docx"
ERD_IMAGE_PATH = TMP_DIR / "erd-thototntu.png"


ACTIVE_TABLES = [
    {
        "name": "users",
        "purpose": "Bang nguoi dung trung tam, luu admin, khach hang va tho.",
        "primary_key": "id",
        "important_columns": [
            "name, email, password",
            "phone, address, avatar",
            "role (admin/customer/worker)",
            "is_active",
        ],
        "relationships": [
            "1-1 voi ho_so_tho theo user_id (chi ap dung cho tho).",
            "1-n voi don_dat_lich theo khach_hang_id.",
            "1-n voi don_dat_lich theo tho_id.",
            "n-n voi danh_muc_dich_vu qua tho_dich_vu.",
            "1-n voi danh_gia o 2 vai tro: nguoi_danh_gia va nguoi_bi_danh_gia.",
            "1-n voi chat_magic theo user_id.",
            "1-n voi app_settings theo updated_by ve mat nghiep vu.",
        ],
        "notes": [
            "Bang trung tam cua he thong.",
            "Role dang luu truc tiep bang enum trong users; bang roles hien khong duoc dung de rang buoc.",
        ],
    },
    {
        "name": "ho_so_tho",
        "purpose": "Ho so bo sung cho user co vai tro tho.",
        "primary_key": "id",
        "important_columns": [
            "user_id",
            "cccd",
            "kinh_nghiem, chung_chi, bang_gia_tham_khao",
            "vi_do, kinh_do",
            "trang_thai_duyet, dang_hoat_dong, trang_thai_hoat_dong",
            "danh_gia_trung_binh, tong_so_danh_gia",
        ],
        "relationships": [
            "n-1 voi users theo user_id.",
        ],
        "notes": [
            "Thuc te nghiep vu la 1 user worker co 1 ho so tho.",
            "Diem danh gia trung binh va tong so danh gia duoc cap nhat tu bang danh_gia.",
        ],
    },
    {
        "name": "danh_muc_dich_vu",
        "purpose": "Danh muc cac dich vu sua chua/bao tri ma he thong cung cap.",
        "primary_key": "id",
        "important_columns": [
            "ten_dich_vu",
            "mo_ta",
            "hinh_anh",
            "trang_thai",
        ],
        "relationships": [
            "n-n voi users qua tho_dich_vu.",
            "n-n voi don_dat_lich qua don_dat_lich_dich_vu.",
        ],
        "notes": [
            "don_dat_lich.dich_vu_id van con ton tai de tuong thich du lieu cu.",
            "Quan he nghiep vu chinh hien nay la qua bang pivot don_dat_lich_dich_vu.",
        ],
    },
    {
        "name": "tho_dich_vu",
        "purpose": "Bang trung gian gan ky nang dich vu cho tho.",
        "primary_key": "id",
        "important_columns": [
            "user_id",
            "dich_vu_id",
        ],
        "relationships": [
            "n-1 voi users theo user_id.",
            "n-1 voi danh_muc_dich_vu theo dich_vu_id.",
        ],
        "notes": [
            "The hien quan he n-n giua tho va dich vu.",
        ],
    },
    {
        "name": "don_dat_lich",
        "purpose": "Bang giao dich chinh cua he thong, luu yeu cau dat lich sua chua.",
        "primary_key": "id",
        "important_columns": [
            "khach_hang_id, tho_id",
            "dich_vu_id (legacy/compatibility)",
            "loai_dat_lich, ngay_hen, khung_gio_hen, thoi_gian_hen",
            "dia_chi, vi_do, kinh_do",
            "mo_ta_van_de, nguyen_nhan, giai_phap",
            "phi_di_lai, phi_linh_kien, tien_cong, tien_thue_xe, tong_tien",
            "phuong_thuc_thanh_toan, trang_thai_thanh_toan",
            "hinh_anh_mo_ta, video_mo_ta, hinh_anh_ket_qua, video_ket_qua",
            "trang_thai, ly_do_huy, thoi_gian_het_han_nhan",
        ],
        "relationships": [
            "n-1 voi users theo khach_hang_id.",
            "n-1 voi users theo tho_id.",
            "n-n voi danh_muc_dich_vu qua don_dat_lich_dich_vu.",
            "1-1 theo nghiep vu voi thanh_toan.",
            "1-1 theo nghiep vu voi danh_gia.",
        ],
        "notes": [
            "Moi don co the gom nhieu dich vu.",
            "thanh_toan va danh_gia dang duoc ep boi business logic, khong chi bang database constraint.",
        ],
    },
    {
        "name": "don_dat_lich_dich_vu",
        "purpose": "Bang pivot gan nhieu dich vu vao mot don dat lich.",
        "primary_key": "id",
        "important_columns": [
            "don_dat_lich_id",
            "dich_vu_id",
            "unique(don_dat_lich_id, dich_vu_id)",
        ],
        "relationships": [
            "n-1 voi don_dat_lich theo don_dat_lich_id.",
            "n-1 voi danh_muc_dich_vu theo dich_vu_id.",
        ],
        "notes": [
            "Bang nay la can cu chinh de mo ta quan he nhieu dich vu trong moi don.",
        ],
    },
    {
        "name": "thanh_toan",
        "purpose": "Luu thong tin giao dich thanh toan cua don dat lich.",
        "primary_key": "id",
        "important_columns": [
            "don_dat_lich_id",
            "so_tien",
            "phuong_thuc (cash/vnpay/momo/zalopay)",
            "ma_giao_dich",
            "trang_thai (pending/success/failed)",
            "thong_tin_extra",
        ],
        "relationships": [
            "n-1 voi don_dat_lich theo don_dat_lich_id.",
        ],
        "notes": [
            "Theo nghiep vu moi don chi thanh toan thanh cong 1 lan.",
            "Code hien chong lap thanh toan success trong PaymentController::processSuccessPayment().",
        ],
    },
    {
        "name": "danh_gia",
        "purpose": "Danh gia cua khach hang sau khi don hoan tat.",
        "primary_key": "id",
        "important_columns": [
            "don_dat_lich_id",
            "nguoi_danh_gia_id",
            "nguoi_bi_danh_gia_id",
            "so_sao, nhan_xet",
            "chuyen_mon, thai_do, dung_gio, gia_ca",
            "so_lan_sua",
        ],
        "relationships": [
            "n-1 voi don_dat_lich theo don_dat_lich_id.",
            "n-1 voi users theo nguoi_danh_gia_id.",
            "n-1 voi users theo nguoi_bi_danh_gia_id.",
        ],
        "notes": [
            "Theo nghiep vu moi don chi co 1 danh gia.",
            "Danh gia duoc sua toi da 1 lan thong qua cot so_lan_sua.",
        ],
    },
    {
        "name": "otp_codes",
        "purpose": "Luu ma OTP xac thuc tam thoi theo email.",
        "primary_key": "id",
        "important_columns": [
            "email",
            "code",
            "expires_at",
        ],
        "relationships": [
            "Khong co khoa ngoai truc tiep.",
        ],
        "notes": [
            "Dung trong dang ky, quen mat khau va xac minh email.",
        ],
    },
    {
        "name": "personal_access_tokens",
        "purpose": "Token dang nhap API cua Laravel Sanctum.",
        "primary_key": "id",
        "important_columns": [
            "tokenable_type, tokenable_id",
            "name, token",
            "abilities, last_used_at, expires_at",
        ],
        "relationships": [
            "Quan he da hinh voi thuc the xac thuc, chu yeu la users.",
        ],
        "notes": [
            "Bang ha tang xac thuc API.",
        ],
    },
    {
        "name": "chat_magic",
        "purpose": "Luu hoi thoai AI/assistant soul phia ung dung.",
        "primary_key": "id",
        "important_columns": [
            "user_id",
            "guest_token",
            "sender",
            "message",
            "meta",
            "created_at",
        ],
        "relationships": [
            "n-1 voi users theo user_id (nullable).",
        ],
        "notes": [
            "Cho phep ca khach chua dang nhap thong qua guest_token.",
        ],
    },
    {
        "name": "chat_messages",
        "purpose": "Bang du phong cho chat message, hien chi co id va timestamps.",
        "primary_key": "id",
        "important_columns": [
            "created_at",
            "updated_at",
        ],
        "relationships": [
            "Chua co khoa ngoai/nghiep vu ro rang trong migration hien tai.",
        ],
        "notes": [
            "Can xem lai neu du an mo rong module chat theo huong rieng.",
        ],
    },
    {
        "name": "app_settings",
        "purpose": "Cau hinh he thong dang key-value.",
        "primary_key": "id",
        "important_columns": [
            "key",
            "value (json)",
            "updated_by",
        ],
        "relationships": [
            "updated_by tham chieu nghiep vu den users, nhung migration hien khong dat foreign key.",
        ],
        "notes": [
            "Dung de luu cau hinh assistant soul va cac cai dat dong.",
        ],
    },
    {
        "name": "roles",
        "purpose": "Bang du phong vai tro, hien migration chi tao id va timestamps.",
        "primary_key": "id",
        "important_columns": [
            "created_at",
            "updated_at",
        ],
        "relationships": [
            "Khong co quan he thuc te trong schema hien tai.",
        ],
        "notes": [
            "He thong hien dang dung users.role thay vi foreign key sang bang roles.",
        ],
    },
]


SYSTEM_TABLES = [
    ("password_reset_tokens", "Luu token dat lai mat khau."),
    ("sessions", "Session web cua Laravel."),
    ("cache", "Du lieu cache key-value."),
    ("cache_locks", "Khoa cache."),
    ("tien_trinh_nen", "Queue jobs."),
    ("lo_tien_trinh_tien", "Queue job batches."),
    ("tien_trinh_that_bai", "Failed jobs."),
]


LEGACY_TABLES = [
    ("bai_dang", "Bang dang bai tim tho theo mo hinh cu. Da bi drop."),
    ("hinh_anh_bai_dang", "Anh cua bai_dang. Da bi drop."),
    ("bao_gia", "Bao gia tu tho cho bai dang. Da bi drop."),
]


ERD_BOXES = [
    {
        "key": "users",
        "title": "users",
        "fields": ["PK id", "name", "email", "role", "is_active"],
        "box": (90, 120, 520, 320),
    },
    {
        "key": "ho_so_tho",
        "title": "ho_so_tho",
        "fields": ["PK id", "FK user_id", "cccd", "trang_thai_duyet", "danh_gia_trung_binh"],
        "box": (90, 430, 520, 650),
    },
    {
        "key": "danh_muc_dich_vu",
        "title": "danh_muc_dich_vu",
        "fields": ["PK id", "ten_dich_vu", "mo_ta", "hinh_anh", "trang_thai"],
        "box": (700, 120, 1130, 320),
    },
    {
        "key": "tho_dich_vu",
        "title": "tho_dich_vu",
        "fields": ["PK id", "FK user_id", "FK dich_vu_id"],
        "box": (700, 430, 1130, 610),
    },
    {
        "key": "don_dat_lich",
        "title": "don_dat_lich",
        "fields": ["PK id", "FK khach_hang_id", "FK tho_id", "legacy dich_vu_id", "trang_thai", "tong_tien"],
        "box": (1320, 120, 1800, 350),
    },
    {
        "key": "don_dat_lich_dich_vu",
        "title": "don_dat_lich_dich_vu",
        "fields": ["PK id", "FK don_dat_lich_id", "FK dich_vu_id", "UNQ booking_service_unique"],
        "box": (1320, 430, 1800, 640),
    },
    {
        "key": "thanh_toan",
        "title": "thanh_toan",
        "fields": ["PK id", "FK don_dat_lich_id", "so_tien", "phuong_thuc", "trang_thai"],
        "box": (1960, 120, 2360, 320),
    },
    {
        "key": "danh_gia",
        "title": "danh_gia",
        "fields": ["PK id", "FK don_dat_lich_id", "FK nguoi_danh_gia_id", "FK nguoi_bi_danh_gia_id", "so_sao", "so_lan_sua"],
        "box": (1960, 430, 2360, 680),
    },
]


ERD_EDGES = [
    ("users", "ho_so_tho", "1-1"),
    ("users", "tho_dich_vu", "1-n"),
    ("danh_muc_dich_vu", "tho_dich_vu", "1-n"),
    ("users", "don_dat_lich", "1-n (khach_hang_id / tho_id)"),
    ("don_dat_lich", "don_dat_lich_dich_vu", "1-n"),
    ("danh_muc_dich_vu", "don_dat_lich_dich_vu", "1-n"),
    ("don_dat_lich", "thanh_toan", "1-1 nghiep vu"),
    ("don_dat_lich", "danh_gia", "1-1 nghiep vu"),
    ("users", "danh_gia", "1-n (reviewer/reviewee)"),
]


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="1a2240", size="10"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def get_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def draw_wrapped(draw, text, xy, width, font, fill):
    words = []
    for paragraph in text.split("\n"):
        if not paragraph:
            words.append("")
            continue
        average_char_width = max(1, int(font.size * 0.55))
        wrapped = wrap(paragraph, max(12, width // average_char_width))
        words.extend(wrapped or [""])
    x, y = xy
    line_height = font.size + 6
    for line in words:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def midpoint_right(box):
    x1, y1, x2, y2 = box
    return x2, (y1 + y2) // 2


def midpoint_left(box):
    x1, y1, x2, y2 = box
    return x1, (y1 + y2) // 2


def midpoint_bottom(box):
    x1, y1, x2, y2 = box
    return (x1 + x2) // 2, y2


def midpoint_top(box):
    x1, y1, x2, y2 = box
    return (x1 + x2) // 2, y1


def connect(draw, src_box, dst_box, label, color):
    sx_r, sy_r = midpoint_right(src_box)
    sx_b, sy_b = midpoint_bottom(src_box)
    dx_l, dy_l = midpoint_left(dst_box)
    dx_t, dy_t = midpoint_top(dst_box)

    if abs(dx_l - sx_r) < abs(dy_t - sy_b):
        start = (sx_r, sy_r)
        end = (dx_l, dy_l)
    else:
        start = (sx_b, sy_b)
        end = (dx_t, dy_t)

    mid_x = (start[0] + end[0]) // 2
    draw.line([start, (mid_x, start[1]), (mid_x, end[1]), end], fill=color, width=5)
    arrow = 12
    draw.polygon(
        [(end[0], end[1]), (end[0] - arrow, end[1] - arrow // 2), (end[0] - arrow, end[1] + arrow // 2)],
        fill=color,
    )
    label_font = get_font(22, bold=True)
    lx = mid_x - 60
    ly = min(start[1], end[1]) + abs(start[1] - end[1]) // 2 - 16
    draw.rounded_rectangle((lx - 10, ly - 6, lx + 170, ly + 34), radius=10, fill="white", outline=color, width=3)
    draw.text((lx, ly), label, font=label_font, fill=color)


def generate_erd_image():
    image = Image.new("RGB", (2480, 1600), "white")
    draw = ImageDraw.Draw(image)
    navy = "#17213f"
    cyan = "#1fb8ff"
    soft = "#eef8ff"
    text = "#111111"

    title_font = get_font(42, bold=True)
    subtitle_font = get_font(24, bold=False)
    box_title_font = get_font(28, bold=True)
    body_font = get_font(21, bold=False)

    draw.text((80, 40), "ERD tong quan database Tho Tot NTU", font=title_font, fill=navy)
    draw.text(
        (80, 92),
        "So do tap trung vao cac bang nghiep vu dang hoat dong. Cac bang he thong duoc mo ta rieng trong tai lieu.",
        font=subtitle_font,
        fill=text,
    )

    key_boxes = {}
    for idx, item in enumerate(ERD_BOXES):
        x1, y1, x2, y2 = item["box"]
        shadow_offset = 12
        draw.rounded_rectangle((x1 + shadow_offset, y1 + shadow_offset, x2 + shadow_offset, y2 + shadow_offset), 22, fill=navy)
        draw.rounded_rectangle((x1, y1, x2, y2), 22, fill="white", outline=navy, width=5)
        draw.rounded_rectangle((x1, y1, x2, y1 + 44), 22, fill=cyan if idx % 2 == 0 else soft, outline=navy, width=5)
        draw.rectangle((x1, y1 + 22, x2, y1 + 44), fill=cyan if idx % 2 == 0 else soft)
        draw.text((x1 + 18, y1 + 8), item["title"], font=box_title_font, fill=navy)
        draw.line((x1, y1 + 56, x2, y1 + 56), fill=navy, width=4)
        y = y1 + 72
        for field in item["fields"]:
            y = draw_wrapped(draw, field, (x1 + 18, y), x2 - x1 - 36, body_font, text)
        key_boxes[item["key"]] = item["box"]

    for src, dst, label in ERD_EDGES:
        connect(draw, key_boxes[src], key_boxes[dst], label, cyan)

    note_box = (80, 1360, 2380, 1530)
    draw.rounded_rectangle((note_box[0] + 10, note_box[1] + 10, note_box[2] + 10, note_box[3] + 10), 20, fill=navy)
    draw.rounded_rectangle(note_box, 20, fill=soft, outline=navy, width=4)
    draw.text((110, 1390), "Ghi chu:", font=get_font(28, bold=True), fill=navy)
    note_text = (
        "1) don_dat_lich.dich_vu_id van con trong schema de tuong thich du lieu cu, "
        "nhung quan he nhieu dich vu hien dang chay qua don_dat_lich_dich_vu. "
        "2) thanh_toan va danh_gia duoc mo ta 1-1 theo nghiep vu hien tai, "
        "du rang buoc chinh dang nam o tang business logic."
    )
    draw_wrapped(draw, note_text, (240, 1388), 2080, get_font(22), text)

    image.save(ERD_IMAGE_PATH)


def add_run(paragraph, text, bold=False, size=11, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=True, size=18 if level == 1 else 14, color="17213F")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=10.5)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "MO TA DATABASE VA SO DO ERD", bold=True, size=22, color="17213F")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(subtitle, "Du an: Website To Find Workers - Tho Tot NTU", size=11)
    subtitle.add_run("\n")
    add_run(subtitle, "Ngay tao tai lieu: 16/03/2026", size=11)

    intro = doc.add_paragraph()
    intro.paragraph_format.space_before = Pt(8)
    intro.paragraph_format.space_after = Pt(8)
    add_run(
        intro,
        "Tai lieu nay tong hop cau truc database hien tai cua du an, mo ta y nghia tung bang, "
        "cac khoa chinh, khoa ngoai, cac quan he giua bang va nhung ghi chu nghiep vu quan trong. "
        "No dong thoi phan biet ro giua bang dang hoat dong, bang he thong va bang legacy da bi loai bo.",
        size=11,
    )

    add_heading(doc, "1. Pham vi va cach doc so do", 1)
    add_bullet(doc, "ERD trong tai lieu tap trung vao cac bang nghiep vu chinh dang duoc su dung.")
    add_bullet(doc, "Cac bang he thong cua Laravel duoc liet ke rieng de tai lieu day du.")
    add_bullet(doc, "Quan he 1-1 cua thanh_toan va danh_gia duoc mo ta theo nghiep vu hien tai.")
    add_bullet(doc, "Cac bang bai_dang, hinh_anh_bai_dang, bao_gia da bi drop va chi duoc nhac o phan legacy.")

    add_heading(doc, "2. So do ERD tong quan", 1)
    doc.add_picture(str(ERD_IMAGE_PATH), width=Inches(9.4))
    erd_caption = doc.add_paragraph()
    erd_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(
        erd_caption,
        "Hinh 1. So do ERD tong quan cho cac bang nghiep vu chinh trong he thong.",
        size=10,
        color="4A4A4A",
    )

    add_heading(doc, "3. Nhom bang nghiep vu chinh", 1)
    for table_meta in ACTIVE_TABLES:
        add_heading(doc, f"3.{ACTIVE_TABLES.index(table_meta) + 1}. Bang {table_meta['name']}", 2)
        p = doc.add_paragraph()
        add_run(p, "Chuc nang: ", bold=True)
        add_run(p, table_meta["purpose"])

        p = doc.add_paragraph()
        add_run(p, "Khoa chinh: ", bold=True)
        add_run(p, table_meta["primary_key"])

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        set_table_borders(table)
        hdr = table.rows[0].cells
        hdr[0].text = "Noi dung"
        hdr[1].text = "Chi tiet"
        for cell in hdr:
            set_cell_shading(cell, "DFF3FF")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
        rows = [
            ("Cot quan trong", "\n".join(f"- {item}" for item in table_meta["important_columns"])),
            ("Quan he", "\n".join(f"- {item}" for item in table_meta["relationships"])),
            ("Ghi chu", "\n".join(f"- {item}" for item in table_meta["notes"])),
        ]
        for left, right in rows:
            row_cells = table.add_row().cells
            row_cells[0].text = left
            row_cells[1].text = right
            row_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            row_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        doc.add_paragraph()

    add_heading(doc, "4. Bang he thong va xac thuc", 1)
    p = doc.add_paragraph()
    add_run(
        p,
        "Day la cac bang phuc vu ha tang Laravel, xac thuc, queue va cache. "
        "Chung khong dong vai tro truc tiep trong ERD nghiep vu, nhung van la mot phan cua database hien tai.",
    )
    for name, desc in SYSTEM_TABLES:
        add_bullet(doc, f"{name}: {desc}")

    add_heading(doc, "5. Bang legacy da bi loai bo", 1)
    p = doc.add_paragraph()
    add_run(
        p,
        "Mot so bang duoc tao o giai doan dau cua du an nhung da bi drop bang migration 2026_03_05_035000. "
        "Cac bang nay khong con la mot phan cua schema hien tai.",
    )
    for name, desc in LEGACY_TABLES:
        add_bullet(doc, f"{name}: {desc}")

    add_heading(doc, "6. Tong hop quan he database", 1)
    relation_lines = [
        "users -> ho_so_tho: 1-1",
        "users -> don_dat_lich (khach_hang_id): 1-n",
        "users -> don_dat_lich (tho_id): 1-n",
        "users <-> danh_muc_dich_vu qua tho_dich_vu: n-n",
        "don_dat_lich <-> danh_muc_dich_vu qua don_dat_lich_dich_vu: n-n",
        "don_dat_lich -> thanh_toan: 1-1 theo nghiep vu",
        "don_dat_lich -> danh_gia: 1-1 theo nghiep vu",
        "danh_gia -> users (nguoi_danh_gia_id / nguoi_bi_danh_gia_id): n-1",
        "chat_magic -> users: n-1, nullable",
        "app_settings -> users(updated_by): lien ket nghiep vu, migration chua dat foreign key",
    ]
    for line in relation_lines:
        add_bullet(doc, line)

    add_heading(doc, "7. Quy tac nghiep vu quan trong", 1)
    business_rules = [
        "Mot don dat lich co the gom nhieu dich vu. Bang pivot chinh la don_dat_lich_dich_vu.",
        "Cot don_dat_lich.dich_vu_id van con ton tai de phuc vu du lieu cu va tuong thich code.",
        "Mot don chi duoc thanh toan thanh cong 1 lan theo business logic trong PaymentController.",
        "Mot don chi co 1 danh gia; danh gia nay co the sua toi da 1 lan thong qua cot so_lan_sua.",
        "Khach hang la chu don; tho co the duoc gan truoc hoac nhan viec sau.",
        "Phi di lai, phi linh kien, tien cong va tien thue xe hop thanh tong_tien cua don.",
    ]
    for rule in business_rules:
        add_bullet(doc, rule)

    add_heading(doc, "8. Van de ky thuat can luu y", 1)
    technical_notes = [
        "Bang roles hien la bang placeholder; he thong thuc te dang dung users.role.",
        "Bang chat_messages hien chi co id va timestamps, chua co du lieu nghiep vu ro rang.",
        "app_settings.updated_by hien chua co foreign key trong migration.",
        "Neu muon rang buoc chat hon, co the bo sung unique constraint cho thanh_toan(don_dat_lich_id) va danh_gia(don_dat_lich_id) neu muon ep 1-1 ngay o tang database.",
    ]
    for note in technical_notes:
        add_bullet(doc, note)

    add_heading(doc, "9. Nguon doi chieu", 1)
    references = [
        r"database/migrations/0001_01_01_000000_create_users_table.php",
        r"database/migrations/2026_02_22_130802_create_ho_so_tho_table.php",
        r"database/migrations/2026_02_22_130747_create_danh_muc_dich_vu_table.php",
        r"database/migrations/2026_02_22_130841_create_don_dat_lich_table.php",
        r"database/migrations/2026_02_22_130842_create_danh_gia_table.php",
        r"database/migrations/2026_03_05_090636_create_thanh_toan_table.php",
        r"database/migrations/2026_03_14_000001_create_don_dat_lich_dich_vu_table.php",
        r"app/Http/Controllers/Api/PaymentController.php",
        r"app/Http/Controllers/Api/DanhGiaController.php",
        r"app/Http/Controllers/Api/DonDatLichController.php",
    ]
    for ref in references:
        add_bullet(doc, ref)

    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    add_heading(doc, "Phu luc A. Danh sach bang hien co", 1)
    appendix = doc.add_table(rows=1, cols=3)
    appendix.alignment = WD_TABLE_ALIGNMENT.CENTER
    appendix.style = "Table Grid"
    set_table_borders(appendix)
    headers = appendix.rows[0].cells
    headers[0].text = "Loai"
    headers[1].text = "Bang"
    headers[2].text = "Ghi chu"
    for cell in headers:
        set_cell_shading(cell, "DFF3FF")
    active_names = {item["name"] for item in ACTIVE_TABLES}
    for name in sorted(active_names):
        cells = appendix.add_row().cells
        cells[0].text = "Nghiep vu/ho tro"
        cells[1].text = name
        cells[2].text = "Dang ton tai"
    for name, desc in SYSTEM_TABLES:
        cells = appendix.add_row().cells
        cells[0].text = "He thong"
        cells[1].text = name
        cells[2].text = desc
    for name, desc in LEGACY_TABLES:
        cells = appendix.add_row().cells
        cells[0].text = "Legacy"
        cells[1].text = name
        cells[2].text = desc

    doc.save(DOC_PATH)


def main():
    ensure_dirs()
    generate_erd_image()
    build_doc()
    print(f"Created DOCX: {DOC_PATH}")
    print(f"Created ERD image: {ERD_IMAGE_PATH}")


if __name__ == "__main__":
    main()
