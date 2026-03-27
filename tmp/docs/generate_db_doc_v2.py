from pathlib import Path
from textwrap import wrap

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"d:\laragon\www\DATN\website-to-find-workers")
OUT_DIR = ROOT / "output" / "doc"
TMP_DIR = ROOT / "tmp" / "docs"
DOC_PATH = OUT_DIR / "mo-ta-database-va-erd-thototntu-chi-tiet.docx"
ERD_CORE_USERS_PATH = TMP_DIR / "erd-core-users-services.png"
ERD_CORE_BOOKING_PATH = TMP_DIR / "erd-core-booking.png"
ERD_SUPPORT_PATH = TMP_DIR / "erd-support-chi-tiet.png"


TABLES = [
    {
        "name": "users",
        "title": "Bảng users",
        "desc": "Lưu thông tin tài khoản người dùng trong hệ thống, gồm quản trị viên, khách hàng và thợ.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("name", "Chuỗi", "VARCHAR(255)", "Họ tên người dùng."),
            ("email", "Chuỗi", "VARCHAR(255)", "Email đăng nhập, duy nhất."),
            ("email_verified_at", "Thời gian", "TIMESTAMP, NULL", "Thời điểm email đã được xác thực."),
            ("password", "Chuỗi", "VARCHAR(255)", "Mật khẩu đã mã hóa."),
            ("phone", "Chuỗi", "VARCHAR(255), NULL", "Số điện thoại liên hệ."),
            ("address", "Chuỗi", "VARCHAR(255), NULL", "Địa chỉ cơ bản của người dùng."),
            ("avatar", "Chuỗi", "VARCHAR(255), NULL", "Đường dẫn ảnh đại diện."),
            ("role", "Liệt kê", "ENUM('admin','customer','worker')", "Vai trò tài khoản."),
            ("is_active", "Boolean", "TINYINT(1)", "Trạng thái hoạt động của tài khoản."),
            ("remember_token", "Chuỗi", "VARCHAR(100), NULL", "Token ghi nhớ đăng nhập."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": [
            "1-1 với bảng ho_so_tho theo user_id.",
            "1-n với bảng don_dat_lich theo khach_hang_id.",
            "1-n với bảng don_dat_lich theo tho_id.",
            "n-n với bảng danh_muc_dich_vu qua bảng trung gian tho_dich_vu.",
            "1-n với bảng danh_gia ở hai vai trò: người đánh giá và người bị đánh giá.",
        ],
    },
    {
        "name": "ho_so_tho",
        "title": "Bảng ho_so_tho",
        "desc": "Lưu hồ sơ mở rộng dành cho tài khoản thợ, bao gồm thông tin xác minh, vị trí và trạng thái làm việc.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("user_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tham chiếu users.id."),
            ("cccd", "Chuỗi", "VARCHAR(255)", "Số CCCD của thợ, duy nhất."),
            ("kinh_nghiem", "Văn bản", "TEXT, NULL", "Mô tả kinh nghiệm làm việc."),
            ("chung_chi", "Chuỗi", "VARCHAR(255), NULL", "Đường dẫn chứng chỉ."),
            ("bang_gia_tham_khao", "Chuỗi", "VARCHAR(255), NULL", "Đường dẫn bảng giá tham khảo."),
            ("vi_do", "Số thập phân", "DECIMAL(10,7), NULL", "Vĩ độ vị trí của thợ."),
            ("kinh_do", "Số thập phân", "DECIMAL(10,7), NULL", "Kinh độ vị trí của thợ."),
            ("ban_kinh_phuc_vu", "Số nguyên", "INT", "Bán kính phục vụ tính theo km."),
            ("trang_thai_duyet", "Liệt kê", "ENUM('cho_duyet','da_duyet','tu_choi')", "Trạng thái duyệt hồ sơ."),
            ("ghi_chu_admin", "Văn bản", "TEXT, NULL", "Ghi chú của quản trị viên."),
            ("dang_hoat_dong", "Boolean", "TINYINT(1)", "Cờ sẵn sàng nhận việc."),
            ("trang_thai_hoat_dong", "Liệt kê", "ENUM('dang_hoat_dong','dang_ban','ngung_hoat_dong','tam_khoa')", "Trạng thái vận hành chi tiết."),
            ("danh_gia_trung_binh", "Số thập phân", "DECIMAL(3,2)", "Điểm đánh giá trung bình của thợ."),
            ("tong_so_danh_gia", "Số nguyên", "INT", "Tổng số đánh giá đã nhận."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": ["n-1 với bảng users theo user_id."],
    },
    {
        "name": "danh_muc_dich_vu",
        "title": "Bảng danh_muc_dich_vu",
        "desc": "Lưu danh sách các dịch vụ mà hệ thống cung cấp cho khách hàng và thợ.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("ten_dich_vu", "Chuỗi", "VARCHAR(255)", "Tên dịch vụ."),
            ("mo_ta", "Chuỗi", "VARCHAR(255), NULL", "Mô tả ngắn cho dịch vụ."),
            ("hinh_anh", "Chuỗi", "VARCHAR(255), NULL", "Ảnh đại diện dịch vụ."),
            ("trang_thai", "Boolean", "TINYINT(1)", "Bật hoặc ẩn dịch vụ."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": [
            "n-n với bảng users qua bảng tho_dich_vu.",
            "n-n với bảng don_dat_lich qua bảng don_dat_lich_dich_vu.",
        ],
    },
    {
        "name": "tho_dich_vu",
        "title": "Bảng tho_dich_vu",
        "desc": "Bảng trung gian dùng để gán các dịch vụ mà mỗi thợ có thể thực hiện.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("user_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tham chiếu users.id."),
            ("dich_vu_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tham chiếu danh_muc_dich_vu.id."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": [
            "n-1 với bảng users theo user_id.",
            "n-1 với bảng danh_muc_dich_vu theo dich_vu_id.",
        ],
    },
    {
        "name": "don_dat_lich",
        "title": "Bảng don_dat_lich",
        "desc": "Lưu toàn bộ thông tin đơn đặt lịch sửa chữa của khách hàng. Đây là bảng trung tâm của luồng nghiệp vụ.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("khach_hang_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tham chiếu users.id của khách hàng."),
            ("tho_id", "Số nguyên lớn", "BIGINT UNSIGNED, NULL", "Khóa ngoại tham chiếu users.id của thợ được gán."),
            ("dich_vu_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Cột tương thích dữ liệu cũ; quan hệ nhiều dịch vụ hiện dùng pivot."),
            ("loai_dat_lich", "Liệt kê", "ENUM('at_home','at_store')", "Hình thức sửa tại nhà hoặc mang đến cửa hàng."),
            ("thoi_gian_hen", "Ngày giờ", "DATETIME", "Thời điểm hẹn tổng quát."),
            ("ngay_hen", "Ngày", "DATE, NULL", "Ngày hẹn theo giao diện wizard."),
            ("khung_gio_hen", "Liệt kê", "ENUM('08:00-10:00','10:00-12:00','12:00-14:00','14:00-17:00'), NULL", "Khung giờ khách chọn."),
            ("khoang_cach", "Số thập phân", "DECIMAL(8,2), NULL", "Khoảng cách từ thợ hoặc cửa hàng tới khách."),
            ("phi_di_lai", "Số thập phân", "DECIMAL(15,2)", "Phí di chuyển dự kiến hoặc thực tế."),
            ("phi_linh_kien", "Số thập phân", "DECIMAL(15,2)", "Phí linh kiện phát sinh."),
            ("thue_xe_cho", "Boolean", "TINYINT(1)", "Khách có yêu cầu thuê xe chở hay không."),
            ("tien_thue_xe", "Số thập phân", "DECIMAL(15,2)", "Chi phí thuê xe chở."),
            ("ghi_chu_linh_kien", "Văn bản", "TEXT, NULL", "Ghi chú về linh kiện hoặc chi phí thêm."),
            ("dia_chi", "Chuỗi", "VARCHAR(255)", "Địa chỉ sửa chữa."),
            ("vi_do", "Số thập phân", "DECIMAL(10,7), NULL", "Vĩ độ vị trí khách hàng."),
            ("kinh_do", "Số thập phân", "DECIMAL(10,7), NULL", "Kinh độ vị trí khách hàng."),
            ("mo_ta_van_de", "Văn bản", "TEXT, NULL", "Mô tả vấn đề do khách nhập."),
            ("nguyen_nhan", "Văn bản", "TEXT, NULL", "Nguyên nhân được AI hoặc hệ thống phân tích."),
            ("giai_phap", "Văn bản", "TEXT, NULL", "Giải pháp đề xuất."),
            ("hinh_anh_mo_ta", "JSON", "JSON, NULL", "Danh sách ảnh mô tả trước khi sửa."),
            ("video_mo_ta", "Chuỗi", "VARCHAR(255), NULL", "Video mô tả trước khi sửa."),
            ("hinh_anh_ket_qua", "JSON", "JSON, NULL", "Danh sách ảnh kết quả sau sửa."),
            ("video_ket_qua", "Chuỗi", "VARCHAR(255), NULL", "Video kết quả sau sửa."),
            ("trang_thai", "Liệt kê", "ENUM('cho_xac_nhan','da_xac_nhan','dang_lam','cho_hoan_thanh','cho_thanh_toan','da_xong','da_huy')", "Trạng thái xử lý đơn."),
            ("tien_cong", "Số thập phân", "DECIMAL(15,2)", "Tiền công sửa chữa."),
            ("ly_do_huy", "Văn bản", "TEXT, NULL", "Lý do hủy đơn nếu có."),
            ("tong_tien", "Số thập phân", "DECIMAL(15,2), NULL", "Tổng chi phí cuối cùng."),
            ("phuong_thuc_thanh_toan", "Liệt kê", "ENUM('cod','transfer')", "Phương thức thu tiền nội bộ của đơn."),
            ("trang_thai_thanh_toan", "Boolean", "TINYINT(1)", "Cờ đã thanh toán hay chưa."),
            ("thoi_gian_het_han_nhan", "Thời gian", "TIMESTAMP, NULL", "Hạn để thợ nhận đơn."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": [
            "n-1 với bảng users theo khach_hang_id.",
            "n-1 với bảng users theo tho_id.",
            "n-n với bảng danh_muc_dich_vu qua bảng don_dat_lich_dich_vu.",
            "1-1 theo nghiệp vụ với bảng thanh_toan.",
            "1-1 theo nghiệp vụ với bảng danh_gia.",
        ],
    },
    {
        "name": "don_dat_lich_dich_vu",
        "title": "Bảng don_dat_lich_dich_vu",
        "desc": "Bảng trung gian thể hiện một đơn đặt lịch có thể chứa nhiều dịch vụ.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("don_dat_lich_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tham chiếu don_dat_lich.id."),
            ("dich_vu_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tham chiếu danh_muc_dich_vu.id."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": [
            "n-1 với bảng don_dat_lich theo don_dat_lich_id.",
            "n-1 với bảng danh_muc_dich_vu theo dich_vu_id.",
        ],
    },
    {
        "name": "thanh_toan",
        "title": "Bảng thanh_toan",
        "desc": "Lưu thông tin giao dịch thanh toán của đơn đặt lịch.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("don_dat_lich_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tham chiếu don_dat_lich.id."),
            ("so_tien", "Số thập phân", "DECIMAL(12,2)", "Số tiền giao dịch."),
            ("phuong_thuc", "Liệt kê", "ENUM('cash','vnpay','momo','zalopay')", "Cổng hoặc cách thanh toán."),
            ("ma_giao_dich", "Chuỗi", "VARCHAR(255), NULL", "Mã giao dịch từ cổng thanh toán."),
            ("trang_thai", "Liệt kê", "ENUM('pending','success','failed')", "Kết quả giao dịch."),
            ("thong_tin_extra", "JSON", "JSON, NULL", "Dữ liệu raw trả về từ gateway."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": ["n-1 với bảng don_dat_lich theo don_dat_lich_id."],
    },
    {
        "name": "danh_gia",
        "title": "Bảng danh_gia",
        "desc": "Lưu đánh giá của khách hàng sau khi đơn đã hoàn tất.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("don_dat_lich_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tham chiếu don_dat_lich.id."),
            ("so_lan_sua", "Số nguyên", "INT", "Số lần chỉnh sửa đánh giá, tối đa 1."),
            ("nguoi_danh_gia_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tới users.id của khách hàng."),
            ("nguoi_bi_danh_gia_id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa ngoại tới users.id của thợ."),
            ("so_sao", "Số nguyên", "INT", "Điểm sao tổng quát từ 1 đến 5."),
            ("nhan_xet", "Văn bản", "TEXT, NULL", "Nhận xét của khách hàng."),
            ("chuyen_mon", "Số nguyên", "INT, NULL", "Điểm về chuyên môn."),
            ("thai_do", "Số nguyên", "INT, NULL", "Điểm về thái độ phục vụ."),
            ("dung_gio", "Số nguyên", "INT, NULL", "Điểm về tính đúng giờ."),
            ("gia_ca", "Số nguyên", "INT, NULL", "Điểm về mức giá."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": [
            "n-1 với bảng don_dat_lich theo don_dat_lich_id.",
            "n-1 với bảng users theo nguoi_danh_gia_id.",
            "n-1 với bảng users theo nguoi_bi_danh_gia_id.",
        ],
    },
    {
        "name": "chat_magic",
        "title": "Bảng chat_magic",
        "desc": "Lưu nội dung trao đổi với trợ lý AI của hệ thống.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("user_id", "Số nguyên lớn", "BIGINT UNSIGNED, NULL", "Khóa ngoại tới users.id, cho phép null."),
            ("guest_token", "Chuỗi", "VARCHAR(100), NULL", "Token khách chưa đăng nhập."),
            ("sender", "Liệt kê", "ENUM('user','assistant','system')", "Nguồn của tin nhắn."),
            ("text", "Văn bản dài", "LONGTEXT", "Nội dung trao đổi."),
            ("meta", "JSON", "JSON, NULL", "Dữ liệu phụ trợ."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
        ],
        "rels": ["n-1 với bảng users theo user_id."],
    },
    {
        "name": "otp_codes",
        "title": "Bảng otp_codes",
        "desc": "Lưu mã OTP tạm thời dùng cho đăng ký, quên mật khẩu và xác minh email.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("email", "Chuỗi", "VARCHAR(255)", "Email nhận OTP."),
            ("code", "Chuỗi", "VARCHAR(6)", "Mã OTP gồm 6 ký tự."),
            ("expires_at", "Thời gian", "TIMESTAMP", "Thời điểm OTP hết hạn."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": ["Không có khóa ngoại trực tiếp."],
    },
    {
        "name": "app_settings",
        "title": "Bảng app_settings",
        "desc": "Lưu cấu hình hệ thống theo dạng khóa - giá trị.",
        "fields": [
            ("id", "Số nguyên lớn", "BIGINT UNSIGNED", "Khóa chính, tự tăng."),
            ("key", "Chuỗi", "VARCHAR(255)", "Tên khóa cấu hình, duy nhất."),
            ("value", "JSON", "JSON, NULL", "Giá trị cấu hình."),
            ("updated_by", "Số nguyên lớn", "BIGINT UNSIGNED, NULL", "ID người cập nhật cuối cùng."),
            ("created_at", "Thời gian", "TIMESTAMP", "Thời điểm tạo bản ghi."),
            ("updated_at", "Thời gian", "TIMESTAMP", "Thời điểm cập nhật gần nhất."),
        ],
        "rels": ["Liên kết nghiệp vụ với bảng users qua updated_by, nhưng migration hiện không khai báo foreign key."],
    },
]


FIELD_LOOKUP = {table["name"]: table["fields"] for table in TABLES}


CORE_USERS_LAYOUT = {
    "canvas": (2600, 1700),
    "boxes": {
        "users": (80, 140, 860, 760),
        "ho_so_tho": (80, 930, 860, 1580),
        "danh_muc_dich_vu": (1350, 140, 2230, 720),
        "tho_dich_vu": (1350, 930, 2230, 1470),
    },
    "relations": [
        ("users", "ho_so_tho", "1 - 1", "cyan"),
        ("users", "tho_dich_vu", "1 - n", "green"),
        ("danh_muc_dich_vu", "tho_dich_vu", "1 - n", "green"),
    ],
}


CORE_BOOKING_LAYOUT = {
    "canvas": (3000, 2000),
    "boxes": {
        "users": (60, 120, 760, 700),
        "danh_muc_dich_vu": (60, 980, 760, 1490),
        "don_dat_lich": (920, 80, 2100, 1580),
        "danh_gia": (2260, 120, 2940, 930),
        "thanh_toan": (2260, 1070, 2940, 1590),
        "don_dat_lich_dich_vu": (920, 1670, 2100, 1920),
    },
    "relations": [
        ("users", "don_dat_lich", "1 - n", "blue"),
        ("users", "danh_gia", "1 - n", "red"),
        ("don_dat_lich", "danh_gia", "1 - 1 NV", "red"),
        ("don_dat_lich", "thanh_toan", "1 - 1 NV", "purple"),
        ("don_dat_lich", "don_dat_lich_dich_vu", "1 - n", "orange"),
        ("danh_muc_dich_vu", "don_dat_lich_dich_vu", "1 - n", "orange"),
    ],
}


SUPPORT_LAYOUT = {
    "canvas": (2800, 1650),
    "boxes": {
        "chat_magic": (80, 150, 860, 900),
        "otp_codes": (980, 150, 1700, 900),
        "app_settings": (1820, 150, 2720, 900),
        "users": (980, 1020, 1700, 1600),
    },
    "relations": [
        ("chat_magic", "users", "n - 1", "blue"),
        ("app_settings", "users", "nghiệp vụ", "purple"),
    ],
}


COLORS = {
    "navy": "#17213F",
    "cyan": "#19B7FF",
    "green": "#2FAE5A",
    "blue": "#2563EB",
    "orange": "#F59E0B",
    "purple": "#8B5CF6",
    "red": "#DC2626",
    "soft": "#F5FBFF",
    "text": "#111827",
}


def ensure_dirs():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    TMP_DIR.mkdir(parents=True, exist_ok=True)


def get_font(size, bold=False):
    candidates = [
        r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf",
        r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_lines(text, font, width_px):
    avg = max(8, int(font.size * 0.52))
    max_chars = max(12, width_px // avg)
    lines = []
    for paragraph in str(text).split("\n"):
        lines.extend(wrap(paragraph, max_chars) or [""])
    return lines


def draw_box(draw, rect, table_name, fields):
    x1, y1, x2, y2 = rect
    shadow = 12
    draw.rounded_rectangle((x1 + shadow, y1 + shadow, x2 + shadow, y2 + shadow), radius=20, fill=COLORS["navy"])
    draw.rounded_rectangle(rect, radius=20, fill="white", outline=COLORS["navy"], width=5)
    draw.rounded_rectangle((x1, y1, x2, y1 + 58), radius=20, fill=COLORS["cyan"], outline=COLORS["navy"], width=5)
    draw.rectangle((x1, y1 + 28, x2, y1 + 58), fill=COLORS["cyan"])
    draw.text((x1 + 18, y1 + 12), table_name, font=get_font(30, True), fill=COLORS["navy"])
    draw.line((x1, y1 + 72, x2, y1 + 72), fill=COLORS["navy"], width=3)

    field_font = get_font(21, False)
    y = y1 + 92
    for name, data_type, _, _ in fields:
        line = f"{name}: {data_type}"
        for wrapped in wrap_lines(line, field_font, x2 - x1 - 36):
            draw.text((x1 + 18, y), wrapped, font=field_font, fill=COLORS["text"])
            y += 29
        y += 2


def mid(rect, side):
    x1, y1, x2, y2 = rect
    if side == "left":
        return (x1, (y1 + y2) // 2)
    if side == "right":
        return (x2, (y1 + y2) // 2)
    if side == "top":
        return ((x1 + x2) // 2, y1)
    return ((x1 + x2) // 2, y2)


def pick_anchors(src_rect, dst_rect):
    sx = (src_rect[0] + src_rect[2]) / 2
    sy = (src_rect[1] + src_rect[3]) / 2
    dx = (dst_rect[0] + dst_rect[2]) / 2
    dy = (dst_rect[1] + dst_rect[3]) / 2
    if abs(dx - sx) > abs(dy - sy):
        return ("right", "left") if dx > sx else ("left", "right")
    return ("bottom", "top") if dy > sy else ("top", "bottom")


def draw_arrow(draw, start, end, color):
    draw.line([start, end], fill=color, width=7)
    arrow = 14
    if abs(end[0] - start[0]) > abs(end[1] - start[1]):
        if end[0] >= start[0]:
            points = [(end[0], end[1]), (end[0] - arrow, end[1] - arrow // 2), (end[0] - arrow, end[1] + arrow // 2)]
        else:
            points = [(end[0], end[1]), (end[0] + arrow, end[1] - arrow // 2), (end[0] + arrow, end[1] + arrow // 2)]
    else:
        if end[1] >= start[1]:
            points = [(end[0], end[1]), (end[0] - arrow // 2, end[1] - arrow), (end[0] + arrow // 2, end[1] - arrow)]
        else:
            points = [(end[0], end[1]), (end[0] - arrow // 2, end[1] + arrow), (end[0] + arrow // 2, end[1] + arrow)]
    draw.polygon(points, fill=color)


def draw_relation(draw, rects, src, dst, label, color_name):
    color = COLORS[color_name]
    src_side, dst_side = pick_anchors(rects[src], rects[dst])
    start = mid(rects[src], src_side)
    end = mid(rects[dst], dst_side)
    if src_side in ("left", "right"):
        mid_x = (start[0] + end[0]) // 2
        points = [start, (mid_x, start[1]), (mid_x, end[1]), end]
    else:
        mid_y = (start[1] + end[1]) // 2
        points = [start, (start[0], mid_y), (end[0], mid_y), end]

    draw.line(points[:-1], fill=color, width=7)
    draw_arrow(draw, points[-2], points[-1], color)

    lx = min(points[1][0], points[2][0]) + abs(points[2][0] - points[1][0]) // 2
    ly = min(points[1][1], points[2][1]) + abs(points[2][1] - points[1][1]) // 2
    label_rect = (lx - 95, ly - 24, lx + 95, ly + 24)
    draw.rounded_rectangle(label_rect, radius=12, fill="white", outline=color, width=4)
    draw.text((label_rect[0] + 18, label_rect[1] + 9), label, font=get_font(20, True), fill=color)


def generate_erd_image(path, layout, title, subtitle):
    image = Image.new("RGB", layout["canvas"], "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 32), title, font=get_font(44, True), fill=COLORS["navy"])
    draw.text((60, 88), subtitle, font=get_font(24, False), fill=COLORS["text"])

    rects = {}
    for table_name, rect in layout["boxes"].items():
        rects[table_name] = rect
        draw_box(draw, rect, table_name, FIELD_LOOKUP[table_name])

    for src, dst, label, color in layout["relations"]:
        draw_relation(draw, rects, src, dst, label, color)

    legend_x = layout["canvas"][0] - 930
    legend_y = 30
    draw.rounded_rectangle((legend_x, legend_y, legend_x + 840, legend_y + 92), radius=18, fill=COLORS["soft"], outline=COLORS["navy"], width=3)
    legend = "Chú thích: 1 - 1 = một - một, 1 - n = một - nhiều, NV = theo nghiệp vụ"
    draw.text((legend_x + 24, legend_y + 28), legend, font=get_font(22, False), fill=COLORS["navy"])
    image.save(path)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table, color="17213F", size="8"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{border_name}"))
        if border is None:
            border = OxmlElement(f"w:{border_name}")
            borders.append(border)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), size)
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), color)


def add_run(paragraph, text, bold=False, size=11, color=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Arial"
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    add_run(p, text, bold=True, size=18 if level == 1 else 14, color="17213F")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    add_run(p, text, size=10.5)


def add_landscape_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)


def add_portrait_section(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def add_attribute_table(doc, table_meta):
    table = doc.add_table(rows=1, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    headers = table.rows[0].cells
    headers[0].text = "STT"
    headers[1].text = "Tên thuộc tính"
    headers[2].text = "Kiểu dữ liệu"
    headers[3].text = "Miền giá trị"
    headers[4].text = "Ghi chú"
    for cell in headers:
        set_cell_shading(cell, "DFF3FF")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
    for idx, field in enumerate(table_meta["fields"], start=1):
        name, data_type, domain, note = field
        row = table.add_row().cells
        row[0].text = str(idx)
        row[1].text = name
        row[2].text = data_type
        row[3].text = domain
        row[4].text = note
        for cell in row:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "MÔ TẢ DATABASE VÀ SƠ ĐỒ ERD CHI TIẾT", True, 21, "17213F")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, "Dự án: Website To Find Workers - Thợ Tốt NTU", size=11)
    p.add_run("\n")
    add_run(p, "Phiên bản chi tiết có tiếng Việt và bảng thuộc tính", size=11)

    intro = doc.add_paragraph()
    add_run(
        intro,
        "Tài liệu này trình bày lại cơ sở dữ liệu theo hướng dễ đọc hơn: sơ đồ ERD rõ mối quan hệ giữa các bảng, "
        "sau đó là phần mô tả chi tiết từng bảng và từng thuộc tính theo đúng mẫu báo cáo thiết kế cơ sở dữ liệu.",
        size=11,
    )

    add_heading(doc, "1. Nguyên tắc đọc tài liệu", 1)
    add_bullet(doc, "Sơ đồ ERD chỉ tập trung vào các bảng nghiệp vụ và bảng hỗ trợ quan trọng.")
    add_bullet(doc, "Các quan hệ 1 - 1 của thanh_toan và danh_gia được hiểu theo nghiệp vụ hiện tại.")
    add_bullet(doc, "Phần mô tả chi tiết dùng cấu trúc: STT - Tên thuộc tính - Kiểu dữ liệu - Miền giá trị - Ghi chú.")
    add_bullet(doc, "Tên bảng và tên cột được giữ nguyên để thuận tiện đối chiếu migration và code.")

    add_landscape_section(doc)
    add_heading(doc, "2. Sơ đồ ERD nhóm người dùng và dịch vụ", 1)
    doc.add_picture(str(ERD_CORE_USERS_PATH), width=Inches(9.6))

    add_landscape_section(doc)
    add_heading(doc, "3. Sơ đồ ERD nhóm đặt lịch, thanh toán và đánh giá", 1)
    doc.add_picture(str(ERD_CORE_BOOKING_PATH), width=Inches(9.6))

    add_landscape_section(doc)
    add_heading(doc, "4. Sơ đồ ERD hỗ trợ", 1)
    doc.add_picture(str(ERD_SUPPORT_PATH), width=Inches(9.6))

    add_portrait_section(doc)
    add_heading(doc, "5. Diễn giải quan hệ giữa các bảng", 1)
    for line in [
        "users - ho_so_tho: Quan hệ một - một. Một tài khoản thợ có một hồ sơ thợ.",
        "users - don_dat_lich: Quan hệ một - nhiều ở hai vai trò. Một khách có nhiều đơn; một thợ có thể được gán nhiều đơn.",
        "users - danh_muc_dich_vu: Quan hệ nhiều - nhiều qua bảng tho_dich_vu.",
        "don_dat_lich - danh_muc_dich_vu: Quan hệ nhiều - nhiều qua bảng don_dat_lich_dich_vu.",
        "don_dat_lich - thanh_toan: Theo nghiệp vụ hiện tại, mỗi đơn thanh toán thành công một lần.",
        "don_dat_lich - danh_gia: Theo nghiệp vụ hiện tại, mỗi đơn có một đánh giá và đánh giá được sửa tối đa một lần.",
        "chat_magic - users: Cho phép lưu hội thoại của user đã đăng nhập hoặc khách bằng guest_token.",
    ]:
        add_bullet(doc, line)

    add_heading(doc, "6. Mô tả chi tiết từng bảng", 1)
    for idx, table_meta in enumerate(TABLES, start=1):
        add_heading(doc, f"6.{idx} {table_meta['title']}", 2)
        p = doc.add_paragraph()
        add_run(p, "Mô tả: ", True)
        add_run(p, table_meta["desc"])
        p = doc.add_paragraph()
        add_run(p, "Quan hệ chính:", True)
        for relation in table_meta["rels"]:
            item = doc.add_paragraph(style="List Bullet")
            add_run(item, relation)
        add_attribute_table(doc, table_meta)

    add_heading(doc, "7. Ghi chú nghiệp vụ quan trọng", 1)
    for note in [
        "Một đơn đặt lịch có thể chứa nhiều dịch vụ; cột dich_vu_id trong don_dat_lich hiện giữ vai trò tương thích dữ liệu cũ.",
        "Một đơn chỉ thanh toán thành công một lần theo business logic trong PaymentController.",
        "Một đơn chỉ có một đánh giá; đánh giá được chỉnh sửa tối đa một lần thông qua cột so_lan_sua.",
        "Bảng app_settings hiện có liên kết nghiệp vụ với users qua updated_by nhưng migration chưa đặt khóa ngoại.",
    ]:
        add_bullet(doc, note)

    doc.save(DOC_PATH)


def main():
    ensure_dirs()
    generate_erd_image(
        ERD_CORE_USERS_PATH,
        CORE_USERS_LAYOUT,
        "Sơ đồ ERD nhóm người dùng và dịch vụ",
        "Tách riêng nhóm bảng này để nhìn rõ các thuộc tính và mũi tên liên kết.",
    )
    generate_erd_image(
        ERD_CORE_BOOKING_PATH,
        CORE_BOOKING_LAYOUT,
        "Sơ đồ ERD nhóm đặt lịch, thanh toán và đánh giá",
        "Tập trung vào bảng đơn đặt lịch và các bảng con có quan hệ trực tiếp với đơn.",
    )
    generate_erd_image(
        ERD_SUPPORT_PATH,
        SUPPORT_LAYOUT,
        "Sơ đồ ERD hỗ trợ - Cấu hình, OTP và hội thoại AI",
        "Nhóm bảng hỗ trợ được tách riêng để sơ đồ chính dễ nhìn hơn.",
    )
    build_doc()
    print(DOC_PATH)


if __name__ == "__main__":
    main()
